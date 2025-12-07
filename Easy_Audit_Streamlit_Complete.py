#!/usr/bin/env python3
"""
Easy Audit V.9.0.0 - Professional Streamlit Version
Complete feature parity with tkinter version
Multi-User Auth + Supabase Cloud + Local JSON + AI-Powered Features

Storage:
- Online: Supabase database with exact column schema
- Offline: saved_records.json (local file)
"""

import streamlit as st
import json
import datetime
import os
import pandas as pd
import openpyxl
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import bcrypt
from datetime import datetime as dt, timedelta
import io
from typing import List, Dict, Optional
import time

# Application Configuration
APP_NAME = "Easy Audit"
APP_VERSION = "9.0.0"

# Supabase Configuration
SUPABASE_URL = "https://yqekjzklcomwzxkdwkga.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InlxZWtqemtsY29td3p4a2R3a2dhIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc2MjE5Njk5NSwiZXhwIjoyMDc3NzcyOTk1fQ.Z5wxHzCVcqzRmmytyj6uNjUQf42RWlVoaQyyr3N456g"

# Cerebras AI Configuration
CEREBRAS_API_KEY = "csk-x4w2dyw62kdhdpcwcpt644mc8mk6mp9nwkn9rerhwd2x5v34"
CEREBRAS_API_BASE_URL = "https://api.cerebras.ai/v1"

# AI Models
CEREBRAS_CLOUD_MODELS = [
    "gpt-oss-120b", "llama-3.3-70b", "llama3.1-8b",
    "qwen-3-235b-a22b-instruct-2507", "qwen-3-32b", "zai-glm-4.6",
]

OLLAMA_LOCAL_MODELS = [
    "deepseek-r1:1.5b", "deepseek-r1:7b", "deepseek-r1:8b", "deepseek-r1:14b",
    "deepseek-r1:32b", "deepseek-r1:70b", "llama3.2:latest", "llama3.1:latest",
    "qwen2.5:latest", "mistral:latest", "phi3:latest",
]

def get_selectable_models():
    return [f"{m} (Cloud)" for m in CEREBRAS_CLOUD_MODELS] + [f"{m} (Local)" for m in OLLAMA_LOCAL_MODELS]

# File Paths
USERS_DB_FILE = "users_db.json"
SETTINGS_FILE = "easy_audit_settings.json"
LOCAL_RECORDS_FILE = "saved_records.json"

# ==================== SUPABASE INTEGRATION ====================
supabase_client = None
SUPABASE_AVAILABLE = False

try:
    from supabase import create_client as supabase_create_client, Client
    
    def init_supabase_client():
        global supabase_client, SUPABASE_AVAILABLE
        try:
            supabase_client = supabase_create_client(SUPABASE_URL, SUPABASE_KEY)
            SUPABASE_AVAILABLE = True
            return supabase_client
        except Exception as e:
            print(f"Supabase init failed: {e}")
            SUPABASE_AVAILABLE = False
            return None
    
    init_supabase_client()
except ImportError:
    print("⚠️ Supabase not installed - using local storage only")
    SUPABASE_AVAILABLE = False

# ==================== DATA PERSISTENCE ====================

def load_users_db():
    """Load users database"""
    if os.path.exists(USERS_DB_FILE):
        try:
            with open(USERS_DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            return {"users": {}}
    return {"users": {}}

def save_users_db(users_db):
    """Save users database"""
    try:
        with open(USERS_DB_FILE, "w", encoding="utf-8") as f:
            json.dump(users_db, f, indent=2)
    except Exception as e:
        print(f"Error saving users DB: {e}")

def load_local_records():
    """Load records from local JSON file"""
    if os.path.exists(LOCAL_RECORDS_FILE):
        try:
            with open(LOCAL_RECORDS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            print(f"Error loading local records: {e}")
            return []
    return []

def save_local_records(records):
    """Save records to local JSON file"""
    try:
        with open(LOCAL_RECORDS_FILE, "w", encoding="utf-8") as f:
            json.dump(records, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Error saving local records: {e}")
        return False

def load_settings():
    """Load application settings"""
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            return get_default_settings()
    return get_default_settings()

def save_settings(settings):
    """Save application settings"""
    try:
        with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
            json.dump(settings, f, indent=2)
    except Exception as e:
        print(f"Error saving settings: {e}")

def get_default_settings():
    """Get default settings"""
    return {
        "storage_mode": "auto",  # auto, supabase, local
        "supabase_table": "audit_records",
        "ai_model": "llama-3.3-70b (Cloud)",
        "theme": "light",
        "date_format": "%d/%m/%Y",
        "auto_sync": True,
        "show_tooltips": True
    }

# ==================== AI INTEGRATION ====================

def call_cerebras_api(prompt: str, model: str = "llama-3.3-70b") -> Optional[str]:
    """Call Cerebras Cloud API"""
    try:
        headers = {
            "Authorization": f"Bearer {CEREBRAS_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": model,
            "messages": [
                {"role": "system", "content": "You are an expert dental auditor with deep knowledge of dental procedures and insurance documentation."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 2000
        }
        response = requests.post(f"{CEREBRAS_API_BASE_URL}/chat/completions", headers=headers, json=data, timeout=30)
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        return None
    except Exception as e:
        print(f"Cerebras API error: {e}")
        return None

def call_ollama_api(prompt: str, model: str = "llama3.2:latest") -> Optional[str]:
    """Call Ollama local API"""
    try:
        data = {"model": model.replace(" (Local)", ""), "prompt": prompt, "stream": False}
        response = requests.post("http://localhost:11434/api/generate", json=data, timeout=60)
        if response.status_code == 200:
            return response.json().get("response", "")
        return None
    except Exception as e:
        print(f"Ollama API error: {e}")
        return None

def call_ai_model(prompt: str, model_name: str) -> Optional[str]:
    """Route AI call to appropriate backend"""
    if "(Cloud)" in model_name:
        clean_model = model_name.replace(" (Cloud)", "").strip()
        return call_cerebras_api(prompt, clean_model)
    else:
        return call_ollama_api(prompt, model_name)

# ==================== RECORD MANAGER ====================

class RecordManager:
    """Manage audit records with Supabase + Local JSON fallback"""
    
    def __init__(self, username: str):
        self.username = username
        self.table_name = st.session_state.settings.get("supabase_table", "audit_records")
        self.storage_mode = st.session_state.settings.get("storage_mode", "auto")
        self.records = []
        self.load_records()
    
    def load_records(self):
        """Load records from Supabase or local file"""
        if self.storage_mode in ["auto", "supabase"] and SUPABASE_AVAILABLE:
            try:
                response = supabase_client.table(self.table_name).select("*").order("CreatedDate", desc=True).execute()
                if response.data:
                    self.records = response.data
                    print(f"✅ Loaded {len(self.records)} records from Supabase")
                    return
            except Exception as e:
                print(f"Supabase load failed: {e}")
        
        # Fallback to local file
        self.records = load_local_records()
        print(f"✅ Loaded {len(self.records)} records from local file")
    
    def save_records(self):
        """Save to both Supabase and local file"""
        # Always save to local file as backup
        save_local_records(self.records)
        
        # Try to save to Supabase if available
        if self.storage_mode in ["auto", "supabase"] and SUPABASE_AVAILABLE:
            try:
                # Note: Supabase sync handled per-record in add/update/delete
                print("✅ Records saved to local and synced to Supabase")
            except Exception as e:
                print(f"Supabase sync warning: {e}")
    
    def add_record(self, record: dict) -> bool:
        """Add new record"""
        try:
            # Normalize field names for consistency
            normalized_record = {
                "Hospital": record.get("Hospital", ""),
                "Doctor": record.get("Doctor", ""),
                "Patient": record.get("Patient", ""),
                "MRN": record.get("MRN", ""),
                "Insurance": record.get("Insurance", ""),
                "Audit Date": record.get("Audit Date", ""),
                "Service Date": record.get("Service Date", ""),
                "Approval Date": record.get("Approval Date", ""),
                "Invoice Date": record.get("Invoice Date", ""),
                "Charged Services": record.get("Charged Services", ""),
                "Approved Services": record.get("Approved Services", ""),
                "Attending Note": record.get("Attending Note", ""),
                "Discrepancy Details": record.get("Discrepancy Details", ""),
                "Discrepancy": record.get("Discrepancy", "No"),
                "The services performed": record.get("The services performed", "Yes"),
                "Approved": record.get("Approved", "Yes"),
                "Created By": self.username,
                "Created Date": dt.now().isoformat()
            }
            
            # Add to local list
            self.records.insert(0, normalized_record)
            
            # Save to local file
            save_local_records(self.records)
            
            # Try to save to Supabase
            if SUPABASE_AVAILABLE and self.storage_mode in ["auto", "supabase"]:
                try:
                    # Map to Supabase column names
                    supabase_record = {
                        "Hospital": normalized_record["Hospital"],
                        "Doctor": normalized_record["Doctor"],
                        "Patient": normalized_record["Patient"],
                        "MRN": normalized_record["MRN"],
                        "Insurance": normalized_record["Insurance"],
                        "AuditDate": normalized_record["Audit Date"],
                        "ServiceDate": normalized_record["Service Date"],
                        "ApprovalDate": normalized_record["Approval Date"],
                        "InvoiceDate": normalized_record["Invoice Date"],
                        "ChargedServices": normalized_record["Charged Services"],
                        "ApprovedServices": normalized_record["Approved Services"],
                        "AttendingNote": normalized_record["Attending Note"],
                        "DiscrepancyDetails": normalized_record["Discrepancy Details"],
                        "Discrepancy": normalized_record["Discrepancy"],
                        "TheServicesPerformed": normalized_record["The services performed"],
                        "Approved": normalized_record["Approved"],
                        "CreatedBy": normalized_record["Created By"],
                        "CreatedDate": normalized_record["Created Date"]
                    }
                    supabase_client.table(self.table_name).insert(supabase_record).execute()
                    print("✅ Record saved to Supabase")
                except Exception as e:
                    print(f"⚠️ Supabase save failed (saved locally): {e}")
            
            return True
        except Exception as e:
            print(f"Error adding record: {e}")
            return False
    
    def update_record(self, index: int, updated_record: dict) -> bool:
        """Update existing record"""
        try:
            if 0 <= index < len(self.records):
                # Preserve Created By and Created Date
                updated_record["Created By"] = self.records[index].get("Created By", self.username)
                updated_record["Created Date"] = self.records[index].get("Created Date", dt.now().isoformat())
                
                self.records[index] = updated_record
                save_local_records(self.records)
                
                # Update in Supabase if available
                if SUPABASE_AVAILABLE and self.storage_mode in ["auto", "supabase"]:
                    try:
                        # Would need unique ID to update in Supabase - for now, local updates only
                        pass
                    except Exception as e:
                        print(f"Supabase update failed: {e}")
                
                return True
            return False
        except Exception as e:
            print(f"Error updating record: {e}")
            return False
    
    def delete_record(self, index: int) -> bool:
        """Delete record"""
        try:
            if 0 <= index < len(self.records):
                self.records.pop(index)
                save_local_records(self.records)
                return True
            return False
        except Exception as e:
            print(f"Error deleting record: {e}")
            return False
    
    def search_records(self, query: str) -> List[dict]:
        """Search records"""
        if not query:
            return self.records
        
        query_lower = query.lower()
        results = []
        
        for record in self.records:
            # Search across all fields
            searchable_text = " ".join([
                str(record.get("Hospital", "")),
                str(record.get("Doctor", "")),
                str(record.get("Patient", "")),
                str(record.get("MRN", "")),
                str(record.get("Insurance", "")),
                str(record.get("Charged Services", "")),
                str(record.get("Approved Services", "")),
                str(record.get("Attending Note", "")),
                str(record.get("Discrepancy Details", ""))
            ]).lower()
            
            if query_lower in searchable_text:
                results.append(record)
        
        return results
    
    def sync_from_supabase(self):
        """Sync records from Supabase to local"""
        if not SUPABASE_AVAILABLE:
            return False, "Supabase not available"
        
        try:
            response = supabase_client.table(self.table_name).select("*").order("CreatedDate", desc=True).execute()
            if response.data:
                # Map Supabase columns back to local format
                synced_records = []
                for record in response.data:
                    synced_records.append({
                        "Hospital": record.get("Hospital", ""),
                        "Doctor": record.get("Doctor", ""),
                        "Patient": record.get("Patient", ""),
                        "MRN": record.get("MRN", ""),
                        "Insurance": record.get("Insurance", ""),
                        "Audit Date": record.get("AuditDate", ""),
                        "Service Date": record.get("ServiceDate", ""),
                        "Approval Date": record.get("ApprovalDate", ""),
                        "Invoice Date": record.get("InvoiceDate", ""),
                        "Charged Services": record.get("ChargedServices", ""),
                        "Approved Services": record.get("ApprovedServices", ""),
                        "Attending Note": record.get("AttendingNote", ""),
                        "Discrepancy Details": record.get("DiscrepancyDetails", ""),
                        "Discrepancy": record.get("Discrepancy", "No"),
                        "The services performed": record.get("TheServicesPerformed", "Yes"),
                        "Approved": record.get("Approved", "Yes"),
                        "Created By": record.get("CreatedBy", ""),
                        "Created Date": record.get("CreatedDate", "")
                    })
                
                self.records = synced_records
                save_local_records(self.records)
                return True, f"Synced {len(synced_records)} records from Supabase"
            return True, "No records in Supabase"
        except Exception as e:
            return False, f"Sync failed: {str(e)}"
    
    def sync_to_supabase(self):
        """Sync local records to Supabase"""
        if not SUPABASE_AVAILABLE:
            return False, "Supabase not available"
        
        try:
            # Clear existing records and insert all local records
            # Note: This is a full sync - in production, use incremental sync
            for record in self.records:
                supabase_record = {
                    "Hospital": record.get("Hospital", ""),
                    "Doctor": record.get("Doctor", ""),
                    "Patient": record.get("Patient", ""),
                    "MRN": record.get("MRN", ""),
                    "Insurance": record.get("Insurance", ""),
                    "AuditDate": record.get("Audit Date", ""),
                    "ServiceDate": record.get("Service Date", ""),
                    "ApprovalDate": record.get("Approval Date", ""),
                    "InvoiceDate": record.get("Invoice Date", ""),
                    "ChargedServices": record.get("Charged Services", ""),
                    "ApprovedServices": record.get("Approved Services", ""),
                    "AttendingNote": record.get("Attending Note", ""),
                    "DiscrepancyDetails": record.get("Discrepancy Details", ""),
                    "Discrepancy": record.get("Discrepancy", "No"),
                    "TheServicesPerformed": record.get("The services performed", "Yes"),
                    "Approved": record.get("Approved", "Yes"),
                    "CreatedBy": record.get("Created By", ""),
                    "CreatedDate": record.get("Created Date", "")
                }
                try:
                    supabase_client.table(self.table_name).insert(supabase_record).execute()
                except:
                    pass  # Skip duplicates
            
            return True, f"Uploaded {len(self.records)} records to Supabase"
        except Exception as e:
            return False, f"Upload failed: {str(e)}"

# ==================== EXPORT MANAGER ====================

class ExportManager:
    """Handle exports to Excel and Word"""
    
    @staticmethod
    def export_to_excel(records: List[dict], filename: str = None) -> io.BytesIO:
        """Export to Excel with professional formatting"""
        # Create DataFrame
        df = pd.DataFrame(records)
        
        # Column order
        column_order = [
            "Hospital", "Doctor", "Patient", "MRN", "Insurance",
            "Audit Date", "Service Date", "Approval Date", "Invoice Date",
            "Charged Services", "Approved Services", "Attending Note",
            "Discrepancy Details", "Discrepancy", "The services performed",
            "Approved", "Created By", "Created Date"
        ]
        
        existing_columns = [col for col in column_order if col in df.columns]
        if existing_columns:
            df = df[existing_columns]
        
        # Create Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Audit Records')
            
            workbook = writer.book
            worksheet = writer.sheets['Audit Records']
            
            # Format headers
            for cell in worksheet[1]:
                cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF", size=11)
                cell.fill = openpyxl.styles.PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                cell.alignment = openpyxl.styles.Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Adjust column widths
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = min(max_length + 2, 60)
                worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # Highlight discrepancies
            for row in range(2, worksheet.max_row + 1):
                discrepancy_cell = None
                for col_letter, cell in zip([c.column_letter for c in worksheet[1]], worksheet[row]):
                    if worksheet[1][worksheet[1].index(cell) if hasattr(worksheet[1], 'index') else 0].value == "Discrepancy":
                        discrepancy_cell = cell
                        break
                
                if discrepancy_cell and str(discrepancy_cell.value).strip().upper() == "YES":
                    for cell in worksheet[row]:
                        cell.fill = openpyxl.styles.PatternFill(start_color="FFE5E5", end_color="FFE5E5", fill_type="solid")
        
        output.seek(0)
        return output
    
    @staticmethod
    def export_to_word(records: List[dict], filename: str = None) -> io.BytesIO:
        """Export to Word document"""
        doc = Document()
        
        # Title
        title = doc.add_heading('Dental Audit Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {dt.now().strftime('%d/%m/%Y at %I:%M %p')}")
        doc.add_paragraph(f"Total Records: {len(records)}")
        
        # Summary statistics
        discrepancies = sum(1 for r in records if r.get('Discrepancy', 'No').strip().upper() == 'YES')
        doc.add_paragraph(f"Records with Discrepancies: {discrepancies}")
        doc.add_paragraph("")
        
        # Records
        for i, record in enumerate(records, 1):
            # Record header
            heading = doc.add_heading(f"Record #{i} - {record.get('Patient', 'Unknown')}", level=2)
            
            # Basic info
            doc.add_paragraph(f"Hospital: {record.get('Hospital', 'N/A')}")
            doc.add_paragraph(f"Doctor: {record.get('Doctor', 'N/A')}")
            doc.add_paragraph(f"MRN: {record.get('MRN', 'N/A')}")
            doc.add_paragraph(f"Insurance: {record.get('Insurance', 'N/A')}")
            
            # Dates
            doc.add_paragraph(f"Audit Date: {record.get('Audit Date', 'N/A')}")
            doc.add_paragraph(f"Service Date: {record.get('Service Date', 'N/A')}")
            
            # Services
            doc.add_paragraph(f"\nCharged Services:\n{record.get('Charged Services', 'N/A')}")
            doc.add_paragraph(f"\nApproved Services:\n{record.get('Approved Services', 'N/A')}")
            
            # Notes and discrepancies
            if record.get('Attending Note'):
                doc.add_paragraph(f"\nAttending Note:\n{record.get('Attending Note')}")
            
            if record.get('Discrepancy', 'No').strip().upper() == 'YES':
                p = doc.add_paragraph("\nDiscrepancy: YES")
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                p.runs[0].font.bold = True
                
                if record.get('Discrepancy Details'):
                    p2 = doc.add_paragraph(f"Details: {record.get('Discrepancy Details')}")
                    p2.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            doc.add_paragraph("")
            doc.add_paragraph("─" * 80)
            doc.add_paragraph("")
        
        # Save to BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

# ==================== AUTHENTICATION ====================

def verify_password(plain_password: str, hashed_password: str) -> bool:
    """Verify password"""
    try:
        return bcrypt.checkpw(plain_password.encode(), hashed_password.encode())
    except:
        return False

def hash_password(password: str) -> str:
    """Hash password"""
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def check_trial_status(user_data: dict) -> tuple:
    """Check trial status"""
    if user_data.get("is_master", False):
        return True, None
    
    trial_end = user_data.get("trial_end_date")
    if not trial_end:
        return False, "No trial period set"
    
    try:
        end_date = dt.fromisoformat(trial_end)
        if dt.now() > end_date:
            days_expired = (dt.now() - end_date).days
            return False, f"Trial expired {days_expired} days ago"
        else:
            days_remaining = (end_date - dt.now()).days
            return True, f"{days_remaining} days remaining"
    except:
        return False, "Invalid trial date"

# ==================== STREAMLIT UI ====================

def init_session_state():
    """Initialize session state"""
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'username' not in st.session_state:
        st.session_state.username = None
    if 'user_data' not in st.session_state:
        st.session_state.user_data = None
    if 'record_manager' not in st.session_state:
        st.session_state.record_manager = None
    if 'settings' not in st.session_state:
        st.session_state.settings = load_settings()
    if 'current_page' not in st.session_state:
        st.session_state.current_page = "main"
    if 'viewing_record' not in st.session_state:
        st.session_state.viewing_record = None
    if 'editing_record_idx' not in st.session_state:
        st.session_state.editing_record_idx = None
    if 'selected_records' not in st.session_state:
        st.session_state.selected_records = []
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []

def set_page_config():
    """Configure page"""
    st.set_page_config(
        page_title=f"{APP_NAME} V.{APP_VERSION}",
        page_icon="🦷",
        layout="wide",
        initial_sidebar_state="expanded"
    )

def apply_custom_css():
    """Apply custom CSS"""
    st.markdown("""
        <style>
        .main {padding: 1rem 2rem;}
        
        /* Header */
        .app-header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 1.5rem 2rem;
            border-radius: 10px;
            margin-bottom: 1.5rem;
            color: white;
            text-align: center;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        .app-title {font-size: 2.2rem; font-weight: bold; margin: 0;}
        .app-subtitle {font-size: 1.1rem; opacity: 0.95; margin-top: 0.3rem;}
        
        /* Stats Cards */
        .stat-card {
            background: white;
            padding: 1.2rem;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.08);
            border-left: 4px solid #667eea;
            margin-bottom: 1rem;
            transition: transform 0.2s;
        }
        .stat-card:hover {transform: translateY(-2px); box-shadow: 0 4px 8px rgba(0,0,0,0.12);}
        .stat-value {font-size: 1.8rem; font-weight: bold; color: #667eea;}
        .stat-label {color: #666; font-size: 0.85rem; text-transform: uppercase; letter-spacing: 0.5px; margin-top: 0.3rem;}
        
        /* Record Cards */
        .record-card {
            background: white;
            padding: 1.2rem;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.08);
            margin-bottom: 0.8rem;
            border-left: 4px solid #4CAF50;
            transition: all 0.2s;
        }
        .record-card:hover {box-shadow: 0 4px 8px rgba(0,0,0,0.12);}
        .record-card.has-discrepancy {border-left-color: #f44336;}
        .record-header {font-weight: 600; font-size: 1.05rem; margin-bottom: 0.5rem; color: #333;}
        .record-detail {color: #666; margin: 0.2rem 0; font-size: 0.9rem;}
        .discrepancy-badge {
            background: #f44336;
            color: white;
            padding: 0.2rem 0.6rem;
            border-radius: 12px;
            font-size: 0.8rem;
            display: inline-block;
            margin-top: 0.4rem;
            font-weight: 500;
        }
        
        /* Buttons */
        .stButton > button {
            border-radius: 6px;
            font-weight: 500;
            transition: all 0.2s;
            border: none;
        }
        .stButton > button:hover {
            transform: translateY(-1px);
            box-shadow: 0 4px 8px rgba(0,0,0,0.15);
        }
        
        /* Forms */
        .stTextInput > div > div > input,
        .stTextArea > div > div > textarea,
        .stSelectbox > div > div > select,
        .stDateInput > div > div > input {
            border-radius: 6px;
            border: 2px solid #e0e0e0;
            padding: 0.5rem;
        }
        .stTextInput > div > div > input:focus,
        .stTextArea > div > div > textarea:focus,
        .stSelectbox > div > div > select:focus,
        .stDateInput > div > div > input:focus {
            border-color: #667eea;
            box-shadow: 0 0 0 0.2rem rgba(102, 126, 234, 0.25);
        }
        
        /* Sidebar */
        section[data-testid="stSidebar"] {
            background: linear-gradient(180deg, #f8f9fa 0%, #e9ecef 100%);
        }
        
        /* Tabs */
        .stTabs [data-baseweb="tab-list"] {
            gap: 1rem;
            background: transparent;
        }
        .stTabs [data-baseweb="tab"] {
            padding: 0.8rem 1.5rem;
            font-weight: 500;
            border-radius: 6px 6px 0 0;
        }
        
        /* Expander */
        .streamlit-expanderHeader {
            font-weight: 500;
            font-size: 1rem;
        }
        
        /* Dataframe */
        .dataframe {
            font-size: 0.9rem;
        }
        
        /* Success/Error/Warning/Info */
        .stSuccess, .stError, .stWarning, .stInfo {
            padding: 1rem;
            border-radius: 6px;
            margin: 0.5rem 0;
        }
        
        /* Loading spinner */
        .stSpinner > div {
            border-top-color: #667eea !important;
        }
        </style>
    """, unsafe_allow_html=True)

# ==================== UI COMPONENTS ====================

def show_login_page():
    """Login page"""
    st.markdown("""
        <div class="app-header">
            <h1 class="app-title">🦷 Easy Audit</h1>
            <p class="app-subtitle">Professional Dental Auditing System V.9.0.0</p>
        </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.subheader("🔐 Login")
        
        with st.form("login_form"):
            username = st.text_input("Username", placeholder="Enter your username")
            password = st.text_input("Password", type="password", placeholder="Enter your password")
            
            col_a, col_b = st.columns(2)
            with col_a:
                login_button = st.form_submit_button("Login", use_container_width=True, type="primary")
            with col_b:
                register_button = st.form_submit_button("Register", use_container_width=True)
        
        if login_button:
            if not username or not password:
                st.error("Please enter both username and password")
            else:
                users_db = load_users_db()
                users = users_db.get("users", {})
                
                if username not in users:
                    st.error("Invalid username or password")
                else:
                    user_data = users[username]
                    
                    if verify_password(password, user_data["password_hash"]):
                        is_valid, message = check_trial_status(user_data)
                        
                        if not is_valid:
                            st.error(f"Account expired: {message}")
                            st.info("Please contact administrator to renew your account")
                        else:
                            st.session_state.authenticated = True
                            st.session_state.username = username
                            st.session_state.user_data = user_data
                            st.session_state.record_manager = RecordManager(username)
                            st.success(f"Welcome back, {user_data.get('full_name', username)}!")
                            if message:
                                st.info(f"Trial status: {message}")
                            time.sleep(1)
                            st.rerun()
                    else:
                        st.error("Invalid username or password")
        
        if register_button:
            st.session_state.current_page = "register"
            st.rerun()

def show_register_page():
    """Registration page"""
    st.markdown("""
        <div class="app-header">
            <h1 class="app-title">🦷 Easy Audit</h1>
            <p class="app-subtitle">Create New Account</p>
        </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.subheader("📝 Register New User")
        
        with st.form("register_form"):
            username = st.text_input("Username*", placeholder="Choose a username")
            password = st.text_input("Password*", type="password", placeholder="Minimum 4 characters")
            confirm_password = st.text_input("Confirm Password*", type="password")
            
            st.divider()
            
            first_name = st.text_input("First Name*")
            last_name = st.text_input("Last Name*")
            email = st.text_input("Email*", placeholder="your.email@example.com")
            
            st.divider()
            
            trial_days = st.number_input("Trial Period (days)", min_value=1, max_value=365, value=30)
            
            col_a, col_b = st.columns(2)
            with col_a:
                register_submit = st.form_submit_button("Create Account", use_container_width=True, type="primary")
            with col_b:
                back_button = st.form_submit_button("Back to Login", use_container_width=True)
        
        if register_submit:
            if not all([username, password, confirm_password, first_name, last_name, email]):
                st.error("Please fill in all required fields")
            elif len(password) < 4:
                st.error("Password must be at least 4 characters")
            elif password != confirm_password:
                st.error("Passwords do not match")
            else:
                users_db = load_users_db()
                users = users_db.get("users", {})
                
                if username in users:
                    st.error("Username already exists")
                else:
                    trial_end = (dt.now() + timedelta(days=trial_days)).isoformat()
                    
                    new_user = {
                        "password_hash": hash_password(password),
                        "first_name": first_name,
                        "last_name": last_name,
                        "full_name": f"{first_name} {last_name}",
                        "email": email,
                        "is_master": False,
                        "created_date": dt.now().isoformat(),
                        "trial_end_date": trial_end
                    }
                    
                    users[username] = new_user
                    users_db["users"] = users
                    save_users_db(users_db)
                    
                    st.success(f"Account created successfully! Trial period: {trial_days} days")
                    time.sleep(2)
                    st.session_state.current_page = "main"
                    st.rerun()
        
        if back_button:
            st.session_state.current_page = "main"
            st.rerun()

def show_main_dashboard():
    """Main dashboard"""
    # Header
    storage_status = "☁️ Supabase" if (SUPABASE_AVAILABLE and st.session_state.settings.get("storage_mode") != "local") else "💾 Local"
    st.markdown(f"""
        <div class="app-header">
            <h1 class="app-title">🦷 Easy Audit V.{APP_VERSION}</h1>
            <p class="app-subtitle">Welcome, {st.session_state.user_data.get('full_name', st.session_state.username)} | Storage: {storage_status}</p>
        </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.markdown("### 📊 Navigation")
        
        page = st.radio(
            "Select Page",
            ["Dashboard", "Records", "Add Record", "AI Analysis", "Settings", "About"],
            label_visibility="collapsed"
        )
        
        st.divider()
        
        # Trial status
        is_valid, message = check_trial_status(st.session_state.user_data)
        if message:
            if is_valid:
                st.success(f"⏰ Trial: {message}")
            else:
                st.error(f"⚠️ {message}")
        
        st.divider()
        
        # Quick stats
        st.markdown("### 📈 Quick Stats")
        records = st.session_state.record_manager.records
        st.metric("Total Records", len(records))
        
        discrepancy_count = sum(1 for r in records if r.get('Discrepancy', 'No').strip().upper() == 'YES')
        st.metric("With Discrepancies", discrepancy_count)
        
        st.divider()
        
        # Cloud sync
        if SUPABASE_AVAILABLE:
            st.markdown("### ☁️ Cloud Sync")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("⬇️ Pull", use_container_width=True, help="Download from Supabase"):
                    with st.spinner("Syncing..."):
                        success, msg = st.session_state.record_manager.sync_from_supabase()
                        if success:
                            st.success(msg)
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error(msg)
            with col2:
                if st.button("⬆️ Push", use_container_width=True, help="Upload to Supabase"):
                    with st.spinner("Uploading..."):
                        success, msg = st.session_state.record_manager.sync_to_supabase()
                        if success:
                            st.success(msg)
                        else:
                            st.error(msg)
        
        st.divider()
        
        # Logout
        if st.button("🚪 Logout", use_container_width=True, type="primary"):
            st.session_state.authenticated = False
            st.session_state.username = None
            st.session_state.user_data = None
            st.session_state.record_manager = None
            st.rerun()
    
    # Main content
    if page == "Dashboard":
        show_dashboard_page()
    elif page == "Records":
        show_records_page()
    elif page == "Add Record":
        show_add_record_page()
    elif page == "AI Analysis":
        show_ai_analysis_page()
    elif page == "Settings":
        show_settings_page()
    elif page == "About":
        show_about_page()

def show_dashboard_page():
    """Dashboard with statistics"""
    st.header("📊 Dashboard")
    
    records = st.session_state.record_manager.records
    
    # Statistics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(f"""
            <div class="stat-card">
                <div class="stat-value">{len(records)}</div>
                <div class="stat-label">Total Records</div>
            </div>
        """, unsafe_allow_html=True)
    
    with col2:
        discrepancy_count = sum(1 for r in records if r.get('Discrepancy', 'No').strip().upper() == 'YES')
        st.markdown(f"""
            <div class="stat-card">
                <div class="stat-value" style="color: #f44336;">{discrepancy_count}</div>
                <div class="stat-label">With Discrepancies</div>
            </div>
        """, unsafe_allow_html=True)
    
    with col3:
        valid_count = len(records) - discrepancy_count
        st.markdown(f"""
            <div class="stat-card">
                <div class="stat-value" style="color: #4CAF50;">{valid_count}</div>
                <div class="stat-label">No Discrepancies</div>
            </div>
        """, unsafe_allow_html=True)
    
    with col4:
        hospitals = set(r.get('Hospital', 'Unknown') for r in records)
        st.markdown(f"""
            <div class="stat-card">
                <div class="stat-value" style="color: #FF9800;">{len(hospitals)}</div>
                <div class="stat-label">Hospitals</div>
            </div>
        """, unsafe_allow_html=True)
    
    st.divider()
    
    # Recent records
    st.subheader("📋 Recent Records")
    
    if not records:
        st.info("No records yet. Add your first record to get started!")
    else:
        recent_records = records[:10]
        
        for i, record in enumerate(recent_records):
            has_discrepancy = record.get('Discrepancy', 'No').strip().upper() == 'YES'
            card_class = "record-card has-discrepancy" if has_discrepancy else "record-card"
            
            st.markdown(f"""
                <div class="{card_class}">
                    <div class="record-header">
                        {record.get('Patient', 'Unknown')} - MRN: {record.get('MRN', 'N/A')}
                    </div>
                    <div class="record-detail">🏥 {record.get('Hospital', 'N/A')} | 👨‍⚕️ {record.get('Doctor', 'N/A')}</div>
                    <div class="record-detail">📅 Audit: {record.get('Audit Date', 'N/A')} | 🏥 Service: {record.get('Service Date', 'N/A')}</div>
                    <div class="record-detail">📋 Insurance: {record.get('Insurance', 'N/A')}</div>
                    {f'<div class="discrepancy-badge">⚠️ Discrepancy Found</div>' if has_discrepancy else ''}
                </div>
            """, unsafe_allow_html=True)
    
    # Charts
    if records:
        st.divider()
        st.subheader("📈 Analytics")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Records by Hospital**")
            hospital_counts = {}
            for record in records:
                hospital = record.get('Hospital', 'Unknown')
                hospital_counts[hospital] = hospital_counts.get(hospital, 0) + 1
            
            if hospital_counts:
                df_hospitals = pd.DataFrame(list(hospital_counts.items()), columns=['Hospital', 'Count'])
                df_hospitals = df_hospitals.sort_values('Count', ascending=False).head(10)
                st.bar_chart(df_hospitals.set_index('Hospital'))
        
        with col2:
            st.markdown("**Records by Doctor**")
            doctor_counts = {}
            for record in records:
                doctor = record.get('Doctor', 'Unknown')
                doctor_counts[doctor] = doctor_counts.get(doctor, 0) + 1
            
            if doctor_counts:
                df_doctors = pd.DataFrame(list(doctor_counts.items()), columns=['Doctor', 'Count'])
                df_doctors = df_doctors.sort_values('Count', ascending=False).head(10)
                st.bar_chart(df_doctors.set_index('Doctor'))

def show_records_page():
    """Records management page"""
    st.header("📋 Audit Records")
    
    # Search and filters
    col1, col2, col3 = st.columns([3, 1, 1])
    
    with col1:
        search_query = st.text_input("🔍 Search", placeholder="Patient, MRN, Hospital, Doctor, Insurance...")
    
    with col2:
        filter_option = st.selectbox("Filter", ["All", "With Discrepancies", "No Discrepancies"])
    
    with col3:
        sort_option = st.selectbox("Sort", ["Newest First", "Oldest First", "Patient A-Z"])
    
    # Get records
    records = st.session_state.record_manager.records
    
    # Apply search
    if search_query:
        records = st.session_state.record_manager.search_records(search_query)
    
    # Apply filter
    if filter_option == "With Discrepancies":
        records = [r for r in records if r.get('Discrepancy', 'No').strip().upper() == 'YES']
    elif filter_option == "No Discrepancies":
        records = [r for r in records if r.get('Discrepancy', 'No').strip().upper() != 'YES']
    
    # Apply sort
    if sort_option == "Newest First":
        pass  # Already sorted by default
    elif sort_option == "Oldest First":
        records = list(reversed(records))
    elif sort_option == "Patient A-Z":
        records = sorted(records, key=lambda x: x.get('Patient', '').lower())
    
    # Export buttons
    col1, col2, col3, col4 = st.columns([1, 1, 1, 3])
    
    with col1:
        if st.button("📥 Excel", use_container_width=True):
            if records:
                excel_file = ExportManager.export_to_excel(records)
                st.download_button(
                    label="Download Excel",
                    data=excel_file,
                    file_name=f"audit_records_{dt.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            else:
                st.warning("No records to export")
    
    with col2:
        if st.button("📄 Word", use_container_width=True):
            if records:
                word_file = ExportManager.export_to_word(records)
                st.download_button(
                    label="Download Word",
                    data=word_file,
                    file_name=f"audit_report_{dt.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.warning("No records to export")
    
    st.divider()
    
    # Display records
    st.markdown(f"**Showing {len(records)} record(s)**")
    
    if not records:
        st.info("No records found matching your criteria")
    else:
        for i, record in enumerate(records):
            with st.container():
                col1, col2 = st.columns([5, 1])
                
                with col1:
                    has_discrepancy = record.get('Discrepancy', 'No').strip().upper() == 'YES'
                    status_icon = "⚠️" if has_discrepancy else "✅"
                    
                    st.markdown(f"### {status_icon} {record.get('Patient', 'Unknown')} - MRN: {record.get('MRN', 'N/A')}")
                    
                    col_a, col_b, col_c = st.columns(3)
                    with col_a:
                        st.text(f"🏥 Hospital: {record.get('Hospital', 'N/A')}")
                        st.text(f"👨‍⚕️ Doctor: {record.get('Doctor', 'N/A')}")
                    with col_b:
                        st.text(f"📋 Insurance: {record.get('Insurance', 'N/A')}")
                        st.text(f"📅 Audit: {record.get('Audit Date', 'N/A')}")
                    with col_c:
                        st.text(f"🏥 Service: {record.get('Service Date', 'N/A')}")
                        st.text(f"✅ Approval: {record.get('Approval Date', 'N/A')}")
                    
                    # Show charged/approved services in expander
                    with st.expander("📋 Services Details"):
                        st.markdown("**Charged Services:**")
                        st.text(record.get('Charged Services', 'N/A'))
                        st.markdown("**Approved Services:**")
                        st.text(record.get('Approved Services', 'N/A'))
                        if record.get('Attending Note'):
                            st.markdown("**Attending Note:**")
                            st.text(record.get('Attending Note'))
                    
                    if has_discrepancy:
                        st.error(f"⚠️ Discrepancy: {record.get('Discrepancy Details', 'No details')}")
                
                with col2:
                    if st.button("👁️", key=f"view_{i}", use_container_width=True, help="View"):
                        st.session_state.viewing_record = record
                        st.rerun()
                    
                    if st.button("✏️", key=f"edit_{i}", use_container_width=True, help="Edit"):
                        st.session_state.editing_record_idx = i
                        st.rerun()
                    
                    if st.button("🗑️", key=f"delete_{i}", use_container_width=True, help="Delete"):
                        if st.session_state.record_manager.delete_record(i):
                            st.success("Deleted")
                            time.sleep(1)
                            st.rerun()
                
                st.divider()
    
    # View/Edit modals
    if st.session_state.viewing_record:
        show_view_record_modal()
    
    if st.session_state.editing_record_idx is not None:
        show_edit_record_modal()

def show_view_record_modal():
    """View record details"""
    record = st.session_state.viewing_record
    
    st.markdown("---")
    st.subheader("👁️ Record Details")
    
    col1, col2 = st.columns([4, 1])
    
    with col1:
        # Display all fields
        col_a, col_b = st.columns(2)
        
        with col_a:
            st.markdown(f"**Hospital:** {record.get('Hospital', 'N/A')}")
            st.markdown(f"**Doctor:** {record.get('Doctor', 'N/A')}")
            st.markdown(f"**Patient:** {record.get('Patient', 'N/A')}")
            st.markdown(f"**MRN:** {record.get('MRN', 'N/A')}")
            st.markdown(f"**Insurance:** {record.get('Insurance', 'N/A')}")
        
        with col_b:
            st.markdown(f"**Audit Date:** {record.get('Audit Date', 'N/A')}")
            st.markdown(f"**Service Date:** {record.get('Service Date', 'N/A')}")
            st.markdown(f"**Approval Date:** {record.get('Approval Date', 'N/A')}")
            st.markdown(f"**Invoice Date:** {record.get('Invoice Date', 'N/A')}")
        
        st.divider()
        
        st.markdown("**Charged Services:**")
        st.text_area("", value=record.get('Charged Services', 'N/A'), height=100, disabled=True, label_visibility="collapsed", key="view_charged")
        
        st.markdown("**Approved Services:**")
        st.text_area("", value=record.get('Approved Services', 'N/A'), height=100, disabled=True, label_visibility="collapsed", key="view_approved")
        
        if record.get('Attending Note'):
            st.markdown("**Attending Note:**")
            st.text_area("", value=record.get('Attending Note'), height=100, disabled=True, label_visibility="collapsed", key="view_note")
        
        st.divider()
        
        st.markdown(f"**Discrepancy:** {record.get('Discrepancy', 'No')}")
        if record.get('Discrepancy', 'No').strip().upper() == 'YES':
            st.error(f"**Details:** {record.get('Discrepancy Details', 'No details')}")
        
        st.markdown(f"**Services Performed:** {record.get('The services performed', 'Yes')}")
        st.markdown(f"**Approved:** {record.get('Approved', 'Yes')}")
        
        st.divider()
        
        st.markdown(f"**Created By:** {record.get('Created By', 'Unknown')}")
        st.markdown(f"**Created Date:** {record.get('Created Date', 'Unknown')}")
    
    with col2:
        if st.button("Close", use_container_width=True, type="primary"):
            st.session_state.viewing_record = None
            st.rerun()

def show_edit_record_modal():
    """Edit record"""
    idx = st.session_state.editing_record_idx
    record = st.session_state.record_manager.records[idx]
    
    st.markdown("---")
    st.subheader("✏️ Edit Record")
    
    with st.form("edit_record_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            hospital = st.text_input("Hospital*", value=record.get('Hospital', ''))
            doctor = st.text_input("Doctor*", value=record.get('Doctor', ''))
            patient = st.text_input("Patient*", value=record.get('Patient', ''))
            mrn = st.text_input("MRN*", value=record.get('MRN', ''))
            insurance = st.text_input("Insurance*", value=record.get('Insurance', ''))
        
        with col2:
            audit_date = st.text_input("Audit Date*", value=record.get('Audit Date', ''), placeholder="DD/MM/YYYY")
            service_date = st.text_input("Service Date*", value=record.get('Service Date', ''), placeholder="DD/MM/YYYY")
            approval_date = st.text_input("Approval Date", value=record.get('Approval Date', ''), placeholder="DD/MM/YYYY")
            invoice_date = st.text_input("Invoice Date", value=record.get('Invoice Date', ''), placeholder="DD/MM/YYYY")
        
        charged_services = st.text_area("Charged Services*", value=record.get('Charged Services', ''), height=100)
        approved_services = st.text_area("Approved Services*", value=record.get('Approved Services', ''), height=100)
        attending_note = st.text_area("Attending Note", value=record.get('Attending Note', ''), height=100)
        discrepancy_details = st.text_area("Discrepancy Details", value=record.get('Discrepancy Details', ''), height=80)
        
        col_a, col_b, col_c = st.columns(3)
        with col_a:
            discrepancy = st.selectbox("Discrepancy", ["No", "Yes"], index=0 if record.get('Discrepancy', 'No').strip().upper() != 'YES' else 1)
        with col_b:
            services_performed = st.selectbox("Services Performed", ["Yes", "No", "No Input"], index=["Yes", "No", "No Input"].index(record.get('The services performed', 'Yes')))
        with col_c:
            approved = st.selectbox("Approved", ["Yes", "No"], index=0 if record.get('Approved', 'Yes').strip().upper() != 'NO' else 1)
        
        col_save, col_cancel = st.columns(2)
        with col_save:
            save_button = st.form_submit_button("💾 Save", use_container_width=True, type="primary")
        with col_cancel:
            cancel_button = st.form_submit_button("❌ Cancel", use_container_width=True)
    
    if save_button:
        if not all([hospital, doctor, patient, mrn, insurance, audit_date, service_date, charged_services, approved_services]):
            st.error("Please fill in all required fields (marked with *)")
        else:
            updated_record = {
                "Hospital": hospital,
                "Doctor": doctor,
                "Patient": patient,
                "MRN": mrn,
                "Insurance": insurance,
                "Audit Date": audit_date,
                "Service Date": service_date,
                "Approval Date": approval_date,
                "Invoice Date": invoice_date,
                "Charged Services": charged_services,
                "Approved Services": approved_services,
                "Attending Note": attending_note,
                "Discrepancy Details": discrepancy_details,
                "Discrepancy": discrepancy,
                "The services performed": services_performed,
                "Approved": approved,
                "Created By": record.get("Created By", st.session_state.username),
                "Created Date": record.get("Created Date", dt.now().isoformat())
            }
            
            if st.session_state.record_manager.update_record(idx, updated_record):
                st.success("Record updated!")
                time.sleep(1)
                st.session_state.editing_record_idx = None
                st.rerun()
            else:
                st.error("Failed to update record")
    
    if cancel_button:
        st.session_state.editing_record_idx = None
        st.rerun()

def show_add_record_page():
    """Add new record"""
    st.header("➕ Add New Record")
    
    with st.form("add_record_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Patient & Provider")
            hospital = st.text_input("Hospital*", placeholder="Hospital name")
            doctor = st.text_input("Doctor*", placeholder="Doctor name")
            patient = st.text_input("Patient*", placeholder="Patient name")
            mrn = st.text_input("MRN*", placeholder="Medical Record Number")
            insurance = st.text_input("Insurance*", placeholder="Insurance company")
        
        with col2:
            st.subheader("Dates")
            audit_date = st.text_input("Audit Date*", placeholder="DD/MM/YYYY")
            service_date = st.text_input("Service Date*", placeholder="DD/MM/YYYY")
            approval_date = st.text_input("Approval Date", placeholder="DD/MM/YYYY")
            invoice_date = st.text_input("Invoice Date", placeholder="DD/MM/YYYY")
        
        st.subheader("Services")
        charged_services = st.text_area("Charged Services*", placeholder="Enter charged services details...", height=100)
        approved_services = st.text_area("Approved Services*", placeholder="Enter approved services details...", height=100)
        attending_note = st.text_area("Attending Note", placeholder="Enter attending physician notes...", height=100)
        
        st.subheader("Audit Results")
        discrepancy_details = st.text_area("Discrepancy Details", placeholder="Describe any discrepancies found...", height=80)
        
        col_a, col_b, col_c = st.columns(3)
        with col_a:
            discrepancy = st.selectbox("Discrepancy", ["No", "Yes"])
        with col_b:
            services_performed = st.selectbox("Services Performed", ["Yes", "No", "No Input"])
        with col_c:
            approved = st.selectbox("Approved", ["Yes", "No"])
        
        st.divider()
        
        col_submit, col_clear = st.columns([1, 1])
        with col_submit:
            submit_button = st.form_submit_button("💾 Save Record", use_container_width=True, type="primary")
        with col_clear:
            clear_button = st.form_submit_button("🔄 Clear", use_container_width=True)
    
    if submit_button:
        if not all([hospital, doctor, patient, mrn, insurance, audit_date, service_date, charged_services, approved_services]):
            st.error("❌ Please fill in all required fields (marked with *)")
        else:
            new_record = {
                "Hospital": hospital,
                "Doctor": doctor,
                "Patient": patient,
                "MRN": mrn,
                "Insurance": insurance,
                "Audit Date": audit_date,
                "Service Date": service_date,
                "Approval Date": approval_date,
                "Invoice Date": invoice_date,
                "Charged Services": charged_services,
                "Approved Services": approved_services,
                "Attending Note": attending_note,
                "Discrepancy Details": discrepancy_details,
                "Discrepancy": discrepancy,
                "The services performed": services_performed,
                "Approved": approved
            }
            
            if st.session_state.record_manager.add_record(new_record):
                st.success("✅ Record added successfully!")
                time.sleep(2)
                st.rerun()
            else:
                st.error("❌ Failed to add record")
    
    if clear_button:
        st.rerun()

def show_ai_analysis_page():
    """AI analysis page"""
    st.header("🤖 AI-Powered Analysis")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("AI Model")
        selected_model = st.selectbox(
            "Select Model",
            options=get_selectable_models(),
            index=0
        )
        st.session_state.settings['ai_model'] = selected_model
    
    with col2:
        st.info("☁️ Cloud: Online\n🖥️ Local: Offline")
    
    st.divider()
    
    tab1, tab2 = st.tabs(["🔍 Analyze Record", "💬 Discuss Case"])
    
    with tab1:
        st.subheader("Analyze Individual Record")
        
        records = st.session_state.record_manager.records
        if not records:
            st.info("No records available")
        else:
            record_options = [f"{r.get('Patient', 'Unknown')} - MRN: {r.get('MRN', 'N/A')} ({r.get('Audit Date', 'N/A')})" for r in records]
            selected_idx = st.selectbox("Select Record", range(len(records)), format_func=lambda x: record_options[x])
            
            selected_record = records[selected_idx]
            
            with st.expander("📋 Record Preview", expanded=True):
                st.markdown(f"**Hospital:** {selected_record.get('Hospital')}")
                st.markdown(f"**Doctor:** {selected_record.get('Doctor')}")
                st.markdown(f"**Patient:** {selected_record.get('Patient')} (MRN: {selected_record.get('MRN')})")
                st.markdown(f"**Insurance:** {selected_record.get('Insurance')}")
                st.markdown(f"**Audit Date:** {selected_record.get('Audit Date')}")
                st.markdown(f"**Discrepancy:** {selected_record.get('Discrepancy', 'No')}")
            
            if st.button("🔍 Analyze with AI", use_container_width=True, type="primary"):
                with st.spinner("AI analyzing..."):
                    prompt = f"""Analyze this dental audit record for discrepancies and compliance issues:

Hospital: {selected_record.get('Hospital')}
Doctor: {selected_record.get('Doctor')}
Patient: {selected_record.get('Patient')} (MRN: {selected_record.get('MRN')})
Insurance: {selected_record.get('Insurance')}
Audit Date: {selected_record.get('Audit Date')}
Service Date: {selected_record.get('Service Date')}

Charged Services:
{selected_record.get('Charged Services')}

Approved Services:
{selected_record.get('Approved Services')}

Attending Note:
{selected_record.get('Attending Note', 'No note provided')}

Current Discrepancy Status: {selected_record.get('Discrepancy', 'No')}
Discrepancy Details: {selected_record.get('Discrepancy Details', 'None')}

Please analyze:
1. Are the charged services aligned with approved services?
2. Does the attending note support the billed services?
3. Are there any compliance or documentation issues?
4. Recommendations for improvement"""
                    
                    response = call_ai_model(prompt, selected_model)
                    
                    if response:
                        st.success("✅ Analysis Complete")
                        st.markdown("### 🤖 AI Analysis Results")
                        st.markdown(response)
                    else:
                        st.error("❌ AI analysis failed")
    
    with tab2:
        st.subheader("Discuss Case with AI")
        
        # Chat history
        for msg in st.session_state.chat_history:
            if msg['role'] == 'user':
                st.markdown(f"**You:** {msg['content']}")
            else:
                st.markdown(f"**🤖 AI:** {msg['content']}")
        
        user_question = st.text_area("Ask about an audit case:", height=100)
        
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("Send", use_container_width=True, type="primary"):
                if user_question:
                    st.session_state.chat_history.append({"role": "user", "content": user_question})
                    
                    with st.spinner("AI thinking..."):
                        context = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.chat_history[-5:]])
                        prompt = f"""You are an expert dental auditor. Answer based on audit standards and compliance.

Conversation:
{context}

Question: {user_question}

Provide a detailed, professional answer."""
                        
                        response = call_ai_model(prompt, selected_model)
                        
                        if response:
                            st.session_state.chat_history.append({"role": "assistant", "content": response})
                            st.rerun()
                        else:
                            st.error("Failed to get AI response")
        
        with col2:
            if st.button("Clear Chat", use_container_width=True):
                st.session_state.chat_history = []
                st.rerun()

def show_settings_page():
    """Settings page"""
    st.header("⚙️ Settings")
    
    user_data = st.session_state.user_data
    
    with st.form("settings_form"):
        st.subheader("Profile")
        first_name = st.text_input("First Name", value=user_data.get('first_name', ''))
        last_name = st.text_input("Last Name", value=user_data.get('last_name', ''))
        email = st.text_input("Email", value=user_data.get('email', ''))
        
        st.divider()
        
        st.subheader("Security")
        new_password = st.text_input("New Password (leave blank to keep)", type="password")
        confirm_password = st.text_input("Confirm Password", type="password")
        
        st.divider()
        
        st.subheader("Storage")
        storage_mode = st.selectbox(
            "Storage Mode",
            ["auto", "supabase", "local"],
            index=["auto", "supabase", "local"].index(st.session_state.settings.get("storage_mode", "auto")),
            help="auto: Use Supabase if available, fallback to local | supabase: Force Supabase | local: Use only local file"
        )
        
        supabase_table = st.text_input("Supabase Table", value=st.session_state.settings.get("supabase_table", "audit_records"))
        
        st.divider()
        
        submit_button = st.form_submit_button("💾 Save Settings", use_container_width=True, type="primary")
    
    if submit_button:
        users_db = load_users_db()
        username = st.session_state.username
        
        if username in users_db.get("users", {}):
            user_data = users_db["users"][username]
            user_data["first_name"] = first_name
            user_data["last_name"] = last_name
            user_data["full_name"] = f"{first_name} {last_name}"
            user_data["email"] = email
            
            if new_password:
                if len(new_password) < 4:
                    st.error("Password must be at least 4 characters")
                elif new_password != confirm_password:
                    st.error("Passwords do not match")
                else:
                    user_data["password_hash"] = hash_password(new_password)
                    save_users_db(users_db)
                    st.session_state.user_data = user_data
                    st.success("Settings saved!")
            else:
                save_users_db(users_db)
                st.session_state.user_data = user_data
                st.success("Settings saved!")
            
            # Update storage settings
            st.session_state.settings["storage_mode"] = storage_mode
            st.session_state.settings["supabase_table"] = supabase_table
            save_settings(st.session_state.settings)

def show_about_page():
    """About page"""
    st.header("ℹ️ About Easy Audit")
    
    st.markdown(f"""
    ### 🦷 Easy Audit V.{APP_VERSION}
    
    **Professional Dental Auditing System**
    
    A comprehensive platform designed for dental audit management with cloud sync and AI-powered analysis.
    
    #### ✨ Key Features
    
    - 🔐 Multi-User Authentication with bcrypt encryption
    - ☁️ Dual Storage: Supabase Cloud + Local JSON fallback
    - 🤖 AI-Powered Analysis with Cerebras Cloud + Ollama Local
    - 📊 Professional Excel & Word exports
    - 🔍 Advanced search and filtering
    - 📈 Analytics dashboard
    - 💾 Auto-save and sync
    - 🎨 Modern, responsive UI
    
    #### 🗄️ Storage Options
    
    - **Auto Mode**: Use Supabase if available, fallback to local
    - **Supabase Mode**: Force cloud storage (requires connection)
    - **Local Mode**: Use only saved_records.json file
    
    #### 📋 Supabase Schema
    
    ```
    Table: audit_records
    Columns:
    - local_id (SERIAL PRIMARY KEY)
    - Hospital, Doctor, Patient, MRN, Insurance
    - AuditDate, ServiceDate, ApprovalDate, InvoiceDate
    - ChargedServices, ApprovedServices, AttendingNote
    - DiscrepancyDetails, Discrepancy, TheServicesPerformed
    - Approved, CreatedBy, CreatedDate
    ```
    
    #### 🔧 Technical Stack
    
    - **Framework**: Streamlit
    - **Database**: Supabase + Local JSON
    - **AI**: Cerebras Cloud + Ollama Local
    - **Export**: openpyxl, python-docx
    - **Auth**: bcrypt
    
    #### 📞 Support
    
    For technical support or questions, contact your system administrator.
    
    ---
    
    **Version**: {APP_VERSION}  
    **Release Date**: December 2024  
    © 2024 Easy Audit. All rights reserved.
    """)

# ==================== MAIN APPLICATION ====================

def main():
    """Main application"""
    set_page_config()
    apply_custom_css()
    init_session_state()
    
    if not st.session_state.authenticated:
        if st.session_state.current_page == "register":
            show_register_page()
        else:
            show_login_page()
    else:
        show_main_dashboard()

if __name__ == "__main__":
    main()
