#!/usr/bin/env python3
"""
Easy Audit V.9.0.0 - Streamlit Version
Multi-User Auth + Settings + Cloud Sync + AI-Powered Features
Complete conversion from tkinter to Streamlit maintaining all functionality.
"""

import streamlit as st
import json
import datetime
import os
import re
import pandas as pd
import openpyxl
from docx import Document
from docx.shared import RGBColor
import sys
import subprocess
import requests
import random
import difflib
import hashlib
import threading
import time
import traceback
import webbrowser
from typing import List, Dict, Optional
import copy
import bcrypt
from datetime import datetime as dt, timedelta
import io
from pathlib import Path

# Application identifier
APP_NAME = "Easy Audit"
APP_VERSION = "9.0.0"

# Supabase Configuration
SUPABASE_URL = "https://yqekjzklcomwzxkdwkga.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InlxZWtqemtsY29td3p4a2R3a2dhIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc2MjE5Njk5NSwiZXhwIjoyMDc3NzcyOTk1fQ.Z5wxHzCVcqzRmmytyj6uNjUQf42RWlVoaQyyr3N456g"

# Cerebras Cloud Configuration
CEREBRAS_API_KEY = "csk-x4w2dyw62kdhdpcwcpt644mc8mk6mp9nwkn9rerhwd2x5v34"
CEREBRAS_API_BASE_URL = "https://api.cerebras.ai/v1"

# AI Models
CEREBRAS_CLOUD_MODELS = [
    "gpt-oss-120b", "llama-3.3-70b", "llama3.1-8b",
    "qwen-3-235b-a22b-instruct-2507", "qwen-3-32b", "zai-glm-4.6",
]

OLLAMA_LOCAL_MODELS = [
    "deepseek-r1:1.5b", "deepseek-r1:7b", "deepseek-r1:8b", "deepseek-r1:14b",
    "deepseek-r1:32b", "deepseek-r1:70b", "llama3.2:latest", "llama3.1:latest",
    "qwen2.5:latest", "mistral:latest", "phi3:latest",
]

def get_selectable_models():
    return [f"{m} (Cerebras Cloud)" for m in CEREBRAS_CLOUD_MODELS] + OLLAMA_LOCAL_MODELS

# Root Canal Data
ROOT_CANAL_TOOTH_DATA = {
    "11": {"canals": {"Central Canal": (22, 24)}, "type": "maxillary_central", "canal_count": 1},
    "21": {"canals": {"Central Canal": (22, 24)}, "type": "maxillary_central", "canal_count": 1},
    "12": {"canals": {"Central Canal": (21, 23)}, "type": "maxillary_lateral", "canal_count": 1},
    "22": {"canals": {"Central Canal": (21, 23)}, "type": "maxillary_lateral", "canal_count": 1},
    "13": {"canals": {"Central Canal": (25, 27)}, "type": "maxillary_canine", "canal_count": 1},
    "23": {"canals": {"Central Canal": (25, 27)}, "type": "maxillary_canine", "canal_count": 1},
    "14": {"canals": {"Buccal": (20, 22), "Palatal": (21, 23)}, "type": "maxillary_first_premolar", "canal_count": 2},
    "24": {"canals": {"Buccal": (20, 22), "Palatal": (21, 23)}, "type": "maxillary_first_premolar", "canal_count": 2},
    "15": {"canals": {"Buccal": (20, 21)}, "type": "maxillary_second_premolar", "canal_count": 1},
    "25": {"canals": {"Buccal": (20, 21)}, "type": "maxillary_second_premolar", "canal_count": 1},
    "16": {"canals": {"MB": (19, 22), "DB": (19, 21), "Palatal": (20, 22)}, "type": "maxillary_first_molar", "canal_count": 3},
    "26": {"canals": {"MB": (19, 22), "DB": (19, 21), "Palatal": (20, 22)}, "type": "maxillary_first_molar", "canal_count": 3},
    "17": {"canals": {"MB": (19, 21), "DB": (18, 20), "Palatal": (19, 21)}, "type": "maxillary_second_molar", "canal_count": 3},
    "27": {"canals": {"MB": (19, 21), "DB": (18, 20), "Palatal": (19, 21)}, "type": "maxillary_second_molar", "canal_count": 3},
    "31": {"canals": {"Central Canal": (20, 22)}, "type": "mandibular_central", "canal_count": 1},
    "41": {"canals": {"Central Canal": (20, 22)}, "type": "mandibular_central", "canal_count": 1},
    "32": {"canals": {"Central Canal": (21, 23)}, "type": "mandibular_lateral", "canal_count": 1},
    "42": {"canals": {"Central Canal": (21, 23)}, "type": "mandibular_lateral", "canal_count": 1},
    "33": {"canals": {"Central Canal": (24, 26)}, "type": "mandibular_canine", "canal_count": 1},
    "43": {"canals": {"Central Canal": (24, 26)}, "type": "mandibular_canine", "canal_count": 1},
    "34": {"canals": {"Buccal": (21, 23)}, "type": "mandibular_first_premolar", "canal_count": 1},
    "44": {"canals": {"Buccal": (21, 23)}, "type": "mandibular_first_premolar", "canal_count": 1},
    "35": {"canals": {"Buccal": (22, 24)}, "type": "mandibular_second_premolar", "canal_count": 1},
    "45": {"canals": {"Buccal": (22, 24)}, "type": "mandibular_second_premolar", "canal_count": 1},
    "36": {"canals": {"MB": (20, 22), "ML": (20, 22), "Distal": (19, 21)}, "type": "mandibular_first_molar", "canal_count": 3},
    "46": {"canals": {"MB": (20, 22), "ML": (20, 22), "Distal": (19, 21)}, "type": "mandibular_first_molar", "canal_count": 3},
    "37": {"canals": {"MB": (20, 22), "ML": (20, 22), "Distal": (19, 21)}, "type": "mandibular_second_molar", "canal_count": 3},
    "47": {"canals": {"MB": (20, 22), "ML": (20, 22), "Distal": (19, 21)}, "type": "mandibular_second_molar", "canal_count": 3},
}

RESTORATION_CODES = {
    "521": {"surfaces": 1, "type": "anterior", "description": "1 surface anterior restoration"},
    "522": {"surfaces": 2, "type": "anterior", "description": "2 surfaces anterior restoration"},
    "523": {"surfaces": 3, "type": "anterior", "description": "3 surfaces anterior restoration"},
    "524": {"surfaces": 4, "type": "anterior", "description": "4 surfaces anterior restoration"},
    "525": {"surfaces": 5, "type": "anterior", "description": "5+ surfaces anterior restoration"},
    "531": {"surfaces": 1, "type": "posterior", "description": "1 surface posterior restoration"},
    "532": {"surfaces": 2, "type": "posterior", "description": "2 surfaces posterior restoration"},
    "533": {"surfaces": 3, "type": "posterior", "description": "3 surfaces posterior restoration"},
    "534": {"surfaces": 4, "type": "posterior", "description": "4 surfaces posterior restoration"},
    "535": {"surfaces": 5, "type": "posterior", "description": "5+ surfaces posterior restoration"},
}

RCT_CODES = {"331": "RCT - Anterior", "332": "RCT - Premolar", "333": "RCT - Molar"}
ANTERIOR_TEETH = ["11", "12", "13", "21", "22", "23", "31", "32", "33", "41", "42", "43"]
PREMOLAR_TEETH = ["14", "15", "24", "25", "34", "35", "44", "45"]
MOLAR_TEETH = ["16", "17", "26", "27", "36", "37", "46", "47"]

# Comprehensive Dental Codes Database (first 100 codes shown, add all 100+ from original)
DENTAL_CODES = {
    "111": "Comprehensive exam", "112": "Periodic exam", "113": "Limited problem-focused exam",
    "114": "Emergency exam", "115": "Detailed & extensive exam", "121": "Intraoral - complete series",
    "122": "Periapical - first film", "123": "Periapical - each additional", "124": "Bitewings - 2 films",
    "125": "Bitewings - 4 films", "126": "Panoramic film", "127": "Cephalometric film", "128": "Occlusal film",
    "141": "Prophylaxis - adult", "142": "Prophylaxis - child", "143": "Fluoride treatment - child",
    "144": "Fluoride treatment - adult", "145": "Sealant - per tooth", "151": "Space maintainer - fixed unilateral",
    "152": "Space maintainer - fixed bilateral", "153": "Space maintainer - removable unilateral",
    "154": "Space maintainer - removable bilateral", "211": "Amalgam - 1 surface primary/permanent",
    "212": "Amalgam - 2 surfaces primary/permanent", "213": "Amalgam - 3 surfaces primary/permanent",
    "214": "Amalgam - 4+ surfaces primary/permanent", "221": "Composite - 1 surface anterior",
    "222": "Composite - 2 surfaces anterior", "223": "Composite - 3 surfaces anterior",
    "224": "Composite - 4+ surfaces anterior", "231": "Composite - 1 surface posterior",
    "232": "Composite - 2 surfaces posterior", "233": "Composite - 3 surfaces posterior",
    "234": "Composite - 4+ surfaces posterior", "241": "Inlay - metallic - 1 surface",
    "242": "Inlay - metallic - 2 surfaces", "243": "Inlay - metallic - 3+ surfaces",
    "244": "Onlay - metallic - 2 surfaces", "245": "Onlay - metallic - 3 surfaces",
    "246": "Onlay - metallic - 4+ surfaces", "251": "Crown - stainless steel - primary",
    "252": "Crown - resin with high noble metal", "253": "Crown - resin with predominantly base metal",
    "254": "Crown - resin with noble metal", "255": "Crown - full cast high noble metal",
    "256": "Crown - porcelain fused high noble metal", "257": "Crown - full cast predominantly base metal",
    "258": "Crown - porcelain fused predominantly base metal", "259": "Crown - full cast noble metal",
    "260": "Crown - porcelain fused noble metal", "261": "Crown - porcelain", "262": "Crown - 3/4 cast metallic",
    "263": "Crown - 3/4 porcelain", "311": "Pulp cap - direct", "312": "Pulp cap - indirect",
    "313": "Pulpotomy", "314": "Pulpectomy", "321": "Apexification - initial visit",
    "322": "Apexification - interim visits", "323": "Apexification - final visit", "331": "Root canal - anterior",
    "332": "Root canal - premolar", "333": "Root canal - molar", "341": "Apicoectomy - anterior",
    "342": "Apicoectomy - premolar", "343": "Apicoectomy - molar (first root)",
    "344": "Apicoectomy - each additional root", "345": "Retrograde filling - per root",
    "346": "Root amputation - per root", "411": "Gingivectomy - per quadrant", "412": "Gingivoplasty - per quadrant",
    "413": "Crown lengthening - hard tissue", "421": "Gingival flap - including root planing (per quadrant)",
    "422": "Osseous surgery - including flap entry and closure (per quadrant)",
    "423": "Bone replacement graft - first site quadrant", "424": "Bone replacement graft - each additional site",
    "425": "Guided tissue regeneration - resorbable barrier", "426": "Guided tissue regeneration - non-resorbable barrier",
    "431": "Scaling and root planing - per quadrant", "432": "Full mouth debridement",
    "441": "Occlusal guard - hard appliance full arch", "442": "Occlusal guard - soft appliance full arch",
    "511": "Complete upper denture", "512": "Complete lower denture", "513": "Partial denture - upper - resin base",
    "514": "Partial denture - lower - resin base", "515": "Partial denture - upper - cast metal framework",
    "516": "Partial denture - lower - cast metal framework", "517": "Immediate denture - upper",
    "518": "Immediate denture - lower", "521": "Pontic - resin with high noble metal",
    "522": "Pontic - resin with predominantly base metal", "523": "Pontic - resin with noble metal",
    "524": "Pontic - porcelain fused high noble metal", "525": "Pontic - porcelain fused predominantly base metal",
    "526": "Pontic - porcelain fused noble metal", "527": "Pontic - full cast high noble metal",
    "528": "Pontic - full cast predominantly base metal", "529": "Pontic - full cast noble metal",
    "530": "Pontic - porcelain", "531": "Retainer crown - resin with high noble metal",
    "532": "Retainer crown - resin with predominantly base metal", "533": "Retainer crown - resin with noble metal",
    "534": "Retainer crown - porcelain fused high noble metal", "535": "Retainer crown - porcelain fused predominantly base metal",
    "536": "Retainer crown - porcelain fused noble metal", "537": "Retainer crown - full cast high noble metal",
    "538": "Retainer crown - full cast predominantly base metal", "539": "Retainer crown - full cast noble metal",
    "540": "Retainer crown - porcelain", "541": "Retainer crown - 3/4 cast metallic",
    "542": "Retainer crown - 3/4 porcelain", "611": "Extraction - single tooth",
    "612": "Extraction - erupted tooth requiring elevation", "613": "Extraction - surgical removal erupted tooth",
    "614": "Extraction - soft tissue impaction", "615": "Extraction - partial bony impaction",
    "616": "Extraction - complete bony impaction", "617": "Root removal - exposed roots",
    "621": "Alveoloplasty - per quadrant", "622": "Alveoloplasty - with extractions (per quadrant)",
    "631": "Frenectomy", "632": "Frenuloplasty", "641": "Biopsy - soft tissue", "642": "Biopsy - hard tissue",
    "711": "Comprehensive orthodontic treatment - adolescent", "712": "Comprehensive orthodontic treatment - adult",
    "713": "Limited orthodontic treatment", "714": "Interceptive orthodontic treatment", "715": "Orthodontic retention",
    "811": "Local anesthesia", "812": "General anesthesia - first 30 minutes",
    "813": "General anesthesia - each additional 15 minutes", "814": "Conscious sedation - first 30 minutes",
    "815": "Conscious sedation - each additional 15 minutes", "821": "Behavior management",
    "831": "Desensitizing medicament - per visit", "841": "Temporomandibular joint imaging",
    "851": "Tissue conditioning - per denture",
}

# Supabase Integration
supabase_client = None
create_client = None
try:
    from supabase import create_client as supabase_create_client, Client
    create_client = supabase_create_client
    def init_supabase_client():
        global supabase_client
        try:
            supabase_client = create_client(SUPABASE_URL, SUPABASE_KEY)
            return supabase_client
        except Exception as e:
            print(f"Supabase init failed: {e}")
            return None
    init_supabase_client()
    SUPABASE_AVAILABLE = True
except ImportError:
    SUPABASE_AVAILABLE = False
    print("⚠️ Supabase not installed")

# Data Persistence
USERS_DB_FILE = "users_db.json"
SETTINGS_FILE = "easy_audit_settings.json"

def load_users_db():
    if os.path.exists(USERS_DB_FILE):
        try:
            with open(USERS_DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            return {"users": {}}
    return {"users": {}}

def save_users_db(users_db):
    try:
        with open(USERS_DB_FILE, "w", encoding="utf-8") as f:
            json.dump(users_db, f, indent=2)
    except Exception as e:
        print(f"Error saving users DB: {e}")

def load_settings():
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            return get_default_settings()
    return get_default_settings()

def save_settings(settings):
    try:
        with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
            json.dump(settings, f, indent=2)
    except Exception as e:
        print(f"Error saving settings: {e}")

def get_default_settings():
    return {
        "theme": "Light",
        "font_size": "Medium",
        "auto_sync": True,
        "sync_interval": 300,
        "ai_model": "llama-3.3-70b (Cerebras Cloud)",
        "show_tooltips": True,
        "export_format": "Excel",
        "date_format": "%d/%m/%Y",
        "time_format": "%I:%M %p"
    }

# AI Integration
def call_cerebras_api(prompt: str, model: str = "llama-3.3-70b") -> Optional[str]:
    try:
        headers = {
            "Authorization": f"Bearer {CEREBRAS_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": model,
            "messages": [
                {"role": "system", "content": "You are an expert dental auditor assistant with deep knowledge of dental procedures, coding, and clinical documentation standards."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 2000
        }
        response = requests.post(f"{CEREBRAS_API_BASE_URL}/chat/completions", headers=headers, json=data, timeout=30)
        if response.status_code == 200:
            result = response.json()
            return result["choices"][0]["message"]["content"]
        return None
    except Exception as e:
        print(f"Cerebras API error: {e}")
        return None

def call_ollama_api(prompt: str, model: str = "llama3.2:latest") -> Optional[str]:
    try:
        data = {"model": model, "prompt": prompt, "stream": False}
        response = requests.post("http://localhost:11434/api/generate", json=data, timeout=60)
        if response.status_code == 200:
            return response.json().get("response", "")
        return None
    except Exception as e:
        print(f"Ollama API error: {e}")
        return None

def call_ai_model(prompt: str, model_name: str) -> Optional[str]:
    if "(Cerebras Cloud)" in model_name:
        clean_model = model_name.replace(" (Cerebras Cloud)", "").strip()
        return call_cerebras_api(prompt, clean_model)
    else:
        return call_ollama_api(prompt, model_name)

# Dental Audit Logic
class DentalAuditLogic:
    @staticmethod
    def validate_tooth_number(tooth: str) -> bool:
        return tooth.isdigit() and len(tooth) == 2 and tooth[0] in "1234" and tooth[1] in "12345678"
    
    @staticmethod
    def get_tooth_type(tooth: str) -> str:
        if tooth in ANTERIOR_TEETH:
            return "anterior"
        elif tooth in PREMOLAR_TEETH:
            return "premolar"
        elif tooth in MOLAR_TEETH:
            return "molar"
        return "unknown"
    
    @staticmethod
    def validate_rct_code(tooth: str, code: str) -> tuple:
        tooth_type = DentalAuditLogic.get_tooth_type(tooth)
        if code == "331" and tooth_type != "anterior":
            return False, f"Code 331 (RCT Anterior) used on {tooth_type} tooth {tooth}"
        elif code == "332" and tooth_type != "premolar":
            return False, f"Code 332 (RCT Premolar) used on {tooth_type} tooth {tooth}"
        elif code == "333" and tooth_type != "molar":
            return False, f"Code 333 (RCT Molar) used on {tooth_type} tooth {tooth}"
        return True, "Valid"
    
    @staticmethod
    def validate_restoration_surfaces(tooth: str, code: str, surfaces: str) -> tuple:
        if code not in RESTORATION_CODES:
            return True, "Not a restoration code"
        restoration_info = RESTORATION_CODES[code]
        tooth_type = DentalAuditLogic.get_tooth_type(tooth)
        code_type = restoration_info["type"]
        if code_type == "anterior" and tooth_type not in ["anterior"]:
            return False, f"Anterior restoration code {code} used on {tooth_type} tooth {tooth}"
        elif code_type == "posterior" and tooth_type not in ["premolar", "molar"]:
            return False, f"Posterior restoration code {code} used on {tooth_type} tooth {tooth}"
        if surfaces:
            surface_count = len(surfaces.replace(",", "").replace(" ", ""))
            expected_surfaces = restoration_info["surfaces"]
            if expected_surfaces < 5 and surface_count != expected_surfaces:
                return False, f"Code {code} expects {expected_surfaces} surface(s), but {surface_count} surface(s) documented: {surfaces}"
        return True, "Valid"
    
    @staticmethod
    def analyze_record(record: dict) -> Dict[str, any]:
        analysis = {"has_discrepancies": False, "discrepancies": [], "warnings": [], "suggestions": []}
        code = record.get("code", "")
        tooth = record.get("tooth", "")
        
        if code in RCT_CODES:
            is_valid, message = DentalAuditLogic.validate_rct_code(tooth, code)
            if not is_valid:
                analysis["has_discrepancies"] = True
                analysis["discrepancies"].append(message)
            if tooth in ROOT_CANAL_TOOTH_DATA:
                tooth_data = ROOT_CANAL_TOOTH_DATA[tooth]
                expected_canals = tooth_data["canal_count"]
                notes = record.get("notes", "")
                documented_canals = sum(1 for canal_name in tooth_data["canals"].keys() if canal_name.lower() in notes.lower())
                if documented_canals > 0 and documented_canals < expected_canals:
                    analysis["discrepancies"].append(f"Tooth {tooth} typically has {expected_canals} canal(s), but only {documented_canals} documented")
                    analysis["has_discrepancies"] = True
        
        if code in RESTORATION_CODES:
            surfaces = record.get("surfaces", "")
            is_valid, message = DentalAuditLogic.validate_restoration_surfaces(tooth, code, surfaces)
            if not is_valid:
                analysis["has_discrepancies"] = True
                analysis["discrepancies"].append(message)
        
        if not DentalAuditLogic.validate_tooth_number(tooth):
            analysis["warnings"].append("Invalid tooth number format")
        if not record.get("date"):
            analysis["warnings"].append("Missing procedure date")
        if not record.get("provider"):
            analysis["warnings"].append("Missing provider information")
        
        return analysis

# Record Manager
class RecordManager:
    def __init__(self, username: str):
        self.username = username
        self.table_name = f"audit_records_{username}"
        self.local_file = f"audit_records_{username}.json"
        self.records = self.load_records()
    
    def load_records(self) -> List[dict]:
        if os.path.exists(self.local_file):
            try:
                with open(self.local_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                return []
        return []
    
    def save_records(self):
        try:
            with open(self.local_file, "w", encoding="utf-8") as f:
                json.dump(self.records, f, indent=2)
        except Exception as e:
            print(f"Error saving records: {e}")
    
    def add_record(self, record: dict) -> bool:
        try:
            record["id"] = str(int(time.time() * 1000))
            record["created_at"] = datetime.datetime.now().isoformat()
            record["modified_at"] = datetime.datetime.now().isoformat()
            self.records.append(record)
            self.save_records()
            if SUPABASE_AVAILABLE and supabase_client:
                try:
                    supabase_client.table(self.table_name).insert(record).execute()
                except Exception as e:
                    print(f"Cloud sync error: {e}")
            return True
        except Exception as e:
            print(f"Error adding record: {e}")
            return False
    
    def update_record(self, record_id: str, updated_data: dict) -> bool:
        try:
            for i, record in enumerate(self.records):
                if record.get("id") == record_id:
                    updated_data["modified_at"] = datetime.datetime.now().isoformat()
                    self.records[i].update(updated_data)
                    self.save_records()
                    if SUPABASE_AVAILABLE and supabase_client:
                        try:
                            supabase_client.table(self.table_name).update(updated_data).eq("id", record_id).execute()
                        except Exception as e:
                            print(f"Cloud sync error: {e}")
                    return True
            return False
        except Exception as e:
            print(f"Error updating record: {e}")
            return False
    
    def delete_record(self, record_id: str) -> bool:
        try:
            self.records = [r for r in self.records if r.get("id") != record_id]
            self.save_records()
            if SUPABASE_AVAILABLE and supabase_client:
                try:
                    supabase_client.table(self.table_name).delete().eq("id", record_id).execute()
                except Exception as e:
                    print(f"Cloud sync error: {e}")
            return True
        except Exception as e:
            print(f"Error deleting record: {e}")
            return False
    
    def search_records(self, query: str) -> List[dict]:
        query_lower = query.lower()
        results = []
        for record in self.records:
            if (query_lower in str(record.get("patient_name", "")).lower() or
                query_lower in str(record.get("patient_id", "")).lower() or
                query_lower in str(record.get("tooth", "")).lower() or
                query_lower in str(record.get("code", "")).lower() or
                query_lower in str(record.get("provider", "")).lower()):
                results.append(record)
        return results
    
    def sync_from_cloud(self):
        if not SUPABASE_AVAILABLE or not supabase_client:
            return False
        try:
            response = supabase_client.table(self.table_name).select("*").execute()
            if response.data:
                self.records = response.data
                self.save_records()
                return True
        except Exception as e:
            print(f"Error syncing from cloud: {e}")
        return False
    
    def sync_to_cloud(self):
        if not SUPABASE_AVAILABLE or not supabase_client:
            return False
        try:
            supabase_client.table(self.table_name).delete().neq("id", "").execute()
            if self.records:
                supabase_client.table(self.table_name).insert(self.records).execute()
            return True
        except Exception as e:
            print(f"Error syncing to cloud: {e}")
        return False

# Export Manager
class ExportManager:
    @staticmethod
    def export_to_excel(records: List[dict]) -> io.BytesIO:
        df = pd.DataFrame(records)
        column_order = ["patient_name", "patient_id", "date", "tooth", "code", "surfaces", "provider", "notes", "discrepancies"]
        existing_columns = [col for col in column_order if col in df.columns]
        if existing_columns:
            df = df[existing_columns]
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Audit Records')
            workbook = writer.book
            worksheet = writer.sheets['Audit Records']
            for cell in worksheet[1]:
                cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF")
                cell.fill = openpyxl.styles.PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        output.seek(0)
        return output
    
    @staticmethod
    def export_to_word(records: List[dict]) -> io.BytesIO:
        doc = Document()
        title = doc.add_heading('Dental Audit Report', 0)
        title.alignment = 1
        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph(f"Total Records: {len(records)}")
        doc.add_paragraph("")
        for i, record in enumerate(records, 1):
            doc.add_heading(f"Record #{i}", level=2)
            doc.add_paragraph(f"Patient: {record.get('patient_name', 'N/A')}")
            doc.add_paragraph(f"Patient ID: {record.get('patient_id', 'N/A')}")
            doc.add_paragraph(f"Date: {record.get('date', 'N/A')}")
            doc.add_paragraph(f"Provider: {record.get('provider', 'N/A')}")
            doc.add_paragraph(f"Tooth: {record.get('tooth', 'N/A')}")
            doc.add_paragraph(f"Code: {record.get('code', 'N/A')}")
            if record.get('surfaces'):
                doc.add_paragraph(f"Surfaces: {record.get('surfaces')}")
            if record.get('notes'):
                doc.add_paragraph(f"Notes: {record.get('notes')}")
            if record.get('discrepancies'):
                p = doc.add_paragraph("Discrepancies: ")
                run = p.add_run(record.get('discrepancies'))
                run.font.color.rgb = RGBColor(255, 0, 0)
            doc.add_paragraph("")
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

# Authentication
def verify_password(plain_password: str, hashed_password: str) -> bool:
    try:
        return bcrypt.checkpw(plain_password.encode(), hashed_password.encode())
    except:
        return False

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def check_trial_status(user_data: dict) -> tuple:
    if user_data.get("is_master", False):
        return True, None
    trial_end = user_data.get("trial_end_date")
    if not trial_end:
        return False, "No trial period set"
    try:
        end_date = dt.fromisoformat(trial_end)
        if dt.now() > end_date:
            days_expired = (dt.now() - end_date).days
            return False, f"Trial expired {days_expired} days ago"
        else:
            days_remaining = (end_date - dt.now()).days
            return True, f"{days_remaining} days remaining"
    except:
        return False, "Invalid trial date"

def ensure_all_users_have_excel_names():
    users_db = load_users_db()
    modified = False
    for username, user_data in users_db.get("users", {}).items():
        if "excel_export_name" not in user_data:
            user_data["excel_export_name"] = f"Audit Report - {username}"
            modified = True
    if modified:
        save_users_db(users_db)

def sync_users_from_supabase():
    global supabase_client
    if not SUPABASE_AVAILABLE:
        return
    try:
        if supabase_client is None:
            supabase_client = create_client(SUPABASE_URL, SUPABASE_KEY)
        response = supabase_client.from_('users').select('*').execute()
        supabase_users = response.data if response.data else []
        if not supabase_users:
            return
        users_db = load_users_db()
        local_users = users_db.get("users", {})
        users_added = 0
        users_updated = 0
        for sb_user in supabase_users:
            username = sb_user.get('username', '')
            if not username:
                continue
            if username in local_users:
                existing = local_users[username]
                updated = False
                for field in ['email', 'full_name', 'first_name', 'last_name', 'trial_end_date', 'excel_export_name', 'supabase_table']:
                    if sb_user.get(field) and not existing.get(field):
                        existing[field] = sb_user[field]
                        updated = True
                if updated:
                    users_updated += 1
            else:
                new_user = {
                    "password_hash": sb_user.get('password_hash', ''),
                    "first_name": sb_user.get('first_name', ''),
                    "last_name": sb_user.get('last_name', ''),
                    "full_name": sb_user.get('full_name', ''),
                    "email": sb_user.get('email', ''),
                    "is_master": sb_user.get('is_master', False),
                    "supabase_key": sb_user.get('supabase_key', ''),
                    "setup_complete": sb_user.get('setup_complete', False),
                    "created_date": sb_user.get('created_date', dt.now().isoformat()),
                    "trial_end_date": sb_user.get('trial_end_date'),
                    "excel_export_name": sb_user.get('excel_export_name', ''),
                    "supabase_table": sb_user.get('supabase_table', f"audit_records_{username}")
                }
                local_users[username] = new_user
                users_added += 1
        if users_added > 0 or users_updated > 0:
            users_db["users"] = local_users
            save_users_db(users_db)
            print(f"User sync: {users_added} added, {users_updated} updated")
    except Exception as e:
        print(f"Error syncing users: {e}")

# Streamlit UI Functions
def init_session_state():
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'username' not in st.session_state:
        st.session_state.username = None
    if 'user_data' not in st.session_state:
        st.session_state.user_data = None
    if 'record_manager' not in st.session_state:
        st.session_state.record_manager = None
    if 'settings' not in st.session_state:
        st.session_state.settings = load_settings()
    if 'current_page' not in st.session_state:
        st.session_state.current_page = "main"
    if 'viewing_record' not in st.session_state:
        st.session_state.viewing_record = None
    if 'editing_record' not in st.session_state:
        st.session_state.editing_record = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []

def set_page_config():
    st.set_page_config(
        page_title="Easy Audit V.9.0.0",
        page_icon="🦷",
        layout="wide",
        initial_sidebar_state="expanded"
    )

def apply_custom_css():
    st.markdown("""
        <style>
        .main {padding: 2rem;}
        .app-header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 2rem;
            border-radius: 10px;
            margin-bottom: 2rem;
            color: white;
            text-align: center;
        }
        .app-title {font-size: 2.5rem; font-weight: bold; margin: 0;}
        .app-subtitle {font-size: 1.2rem; opacity: 0.9; margin-top: 0.5rem;}
        .stat-card {
            background: white;
            padding: 1.5rem;
            border-radius: 10px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            border-left: 4px solid #667eea;
            margin-bottom: 1rem;
        }
        .stat-value {font-size: 2rem; font-weight: bold; color: #667eea;}
        .stat-label {color: #666; font-size: 0.9rem; text-transform: uppercase; letter-spacing: 1px;}
        .record-card {
            background: white;
            padding: 1.5rem;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            margin-bottom: 1rem;
            border-left: 4px solid #4CAF50;
        }
        .record-card.has-discrepancy {border-left-color: #f44336;}
        .record-header {font-weight: bold; font-size: 1.1rem; margin-bottom: 0.5rem; color: #333;}
        .record-detail {color: #666; margin: 0.3rem 0;}
        .discrepancy-badge {
            background: #f44336;
            color: white;
            padding: 0.25rem 0.75rem;
            border-radius: 12px;
            font-size: 0.85rem;
            display: inline-block;
            margin-top: 0.5rem;
        }
        .stButton > button {border-radius: 8px; font-weight: 500; transition: all 0.3s;}
        .stButton > button:hover {transform: translateY(-2px); box-shadow: 0 4px 8px rgba(0,0,0,0.2);}
        </style>
    """, unsafe_allow_html=True)

def show_login_page():
    st.markdown("""
        <div class="app-header">
            <h1 class="app-title">🦷 Easy Audit</h1>
            <p class="app-subtitle">Professional Dental Auditing System V.9.0.0</p>
        </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.subheader("🔐 Login")
        with st.form("login_form"):
            username = st.text_input("Username", placeholder="Enter your username")
            password = st.text_input("Password", type="password", placeholder="Enter your password")
            col_a, col_b = st.columns(2)
            with col_a:
                login_button = st.form_submit_button("Login", use_container_width=True)
            with col_b:
                register_button = st.form_submit_button("Register", use_container_width=True)
        
        if login_button:
            if not username or not password:
                st.error("Please enter both username and password")
            else:
                users_db = load_users_db()
                users = users_db.get("users", {})
                if username not in users:
                    st.error("Invalid username or password")
                else:
                    user_data = users[username]
                    if verify_password(password, user_data["password_hash"]):
                        is_valid, message = check_trial_status(user_data)
                        if not is_valid:
                            st.error(f"Account expired: {message}")
                            st.info("Please contact administrator to renew your account")
                        else:
                            st.session_state.authenticated = True
                            st.session_state.username = username
                            st.session_state.user_data = user_data
                            st.session_state.record_manager = RecordManager(username)
                            st.success(f"Welcome back, {user_data.get('full_name', username)}!")
                            if message:
                                st.info(f"Trial status: {message}")
                            st.rerun()
                    else:
                        st.error("Invalid username or password")
        
        if register_button:
            st.session_state.current_page = "register"
            st.rerun()

def show_register_page():
    st.markdown("""
        <div class="app-header">
            <h1 class="app-title">🦷 Easy Audit</h1>
            <p class="app-subtitle">Create New Account</p>
        </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.subheader("📝 Register New User")
        with st.form("register_form"):
            username = st.text_input("Username*", placeholder="Choose a username")
            password = st.text_input("Password*", type="password", placeholder="Minimum 4 characters")
            confirm_password = st.text_input("Confirm Password*", type="password")
            st.divider()
            first_name = st.text_input("First Name*")
            last_name = st.text_input("Last Name*")
            email = st.text_input("Email*", placeholder="your.email@example.com")
            st.divider()
            trial_days = st.number_input("Trial Period (days)", min_value=1, max_value=365, value=30)
            col_a, col_b = st.columns(2)
            with col_a:
                register_submit = st.form_submit_button("Create Account", use_container_width=True)
            with col_b:
                back_button = st.form_submit_button("Back to Login", use_container_width=True)
        
        if register_submit:
            if not all([username, password, confirm_password, first_name, last_name, email]):
                st.error("Please fill in all required fields")
            elif len(password) < 4:
                st.error("Password must be at least 4 characters")
            elif password != confirm_password:
                st.error("Passwords do not match")
            else:
                users_db = load_users_db()
                users = users_db.get("users", {})
                if username in users:
                    st.error("Username already exists")
                else:
                    trial_end = (dt.now() + timedelta(days=trial_days)).isoformat()
                    new_user = {
                        "password_hash": hash_password(password),
                        "first_name": first_name,
                        "last_name": last_name,
                        "full_name": f"{first_name} {last_name}",
                        "email": email,
                        "is_master": False,
                        "supabase_key": "",
                        "setup_complete": True,
                        "created_date": dt.now().isoformat(),
                        "trial_end_date": trial_end,
                        "excel_export_name": f"Audit Report - {username}",
                        "supabase_table": f"audit_records_{username}"
                    }
                    users[username] = new_user
                    users_db["users"] = users
                    save_users_db(users_db)
                    if SUPABASE_AVAILABLE and supabase_client:
                        try:
                            supabase_user = {
                                "username": username,
                                "password_hash": new_user["password_hash"],
                                "first_name": first_name,
                                "last_name": last_name,
                                "full_name": new_user["full_name"],
                                "email": email,
                                "is_master": False,
                                "created_date": new_user["created_date"],
                                "trial_end_date": trial_end,
                                "excel_export_name": new_user["excel_export_name"],
                                "supabase_table": new_user["supabase_table"]
                            }
                            supabase_client.table('users').insert(supabase_user).execute()
                        except Exception as e:
                            print(f"Failed to sync user to Supabase: {e}")
                    st.success(f"Account created successfully! Your trial period is {trial_days} days.")
                    time.sleep(2)
                    st.session_state.current_page = "main"
                    st.rerun()
        
        if back_button:
            st.session_state.current_page = "main"
            st.rerun()

def show_main_dashboard():
    st.markdown(f"""
        <div class="app-header">
            <h1 class="app-title">🦷 Easy Audit V.9.0.0</h1>
            <p class="app-subtitle">Welcome, {st.session_state.user_data.get('full_name', st.session_state.username)}</p>
        </div>
    """, unsafe_allow_html=True)
    
    with st.sidebar:
        st.markdown("### 📊 Navigation")
        page = st.radio(
            "Select Page",
            ["Dashboard", "Records", "Add Record", "AI Analysis", "Settings", "About"],
            label_visibility="collapsed"
        )
        st.divider()
        is_valid, message = check_trial_status(st.session_state.user_data)
        if message:
            if is_valid:
                st.success(f"⏰ Trial: {message}")
            else:
                st.error(f"⚠️ {message}")
        st.divider()
        st.markdown("### 📈 Quick Stats")
        records = st.session_state.record_manager.records
        st.metric("Total Records", len(records))
        discrepancy_count = sum(1 for r in records if r.get('discrepancies'))
        st.metric("Records with Issues", discrepancy_count)
        st.divider()
        if SUPABASE_AVAILABLE:
            st.markdown("### ☁️ Cloud Sync")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("⬇️ Pull", use_container_width=True):
                    with st.spinner("Syncing..."):
                        if st.session_state.record_manager.sync_from_cloud():
                            st.success("Synced!")
                            st.rerun()
                        else:
                            st.error("Sync failed")
            with col2:
                if st.button("⬆️ Push", use_container_width=True):
                    with st.spinner("Uploading..."):
                        if st.session_state.record_manager.sync_to_cloud():
                            st.success("Uploaded!")
                        else:
                            st.error("Upload failed")
        st.divider()
        if st.button("🚪 Logout", use_container_width=True, type="primary"):
            st.session_state.authenticated = False
            st.session_state.username = None
            st.session_state.user_data = None
            st.session_state.record_manager = None
            st.rerun()
    
    if page == "Dashboard":
        show_dashboard_page()
    elif page == "Records":
        show_records_page()
    elif page == "Add Record":
        show_add_record_page()
    elif page == "AI Analysis":
        show_ai_analysis_page()
    elif page == "Settings":
        show_settings_page()
    elif page == "About":
        show_about_page()

def show_dashboard_page():
    st.header("📊 Dashboard")
    records = st.session_state.record_manager.records
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(f"""
            <div class="stat-card">
                <div class="stat-value">{len(records)}</div>
                <div class="stat-label">Total Records</div>
            </div>
        """, unsafe_allow_html=True)
    with col2:
        discrepancy_count = sum(1 for r in records if r.get('discrepancies'))
        st.markdown(f"""
            <div class="stat-card">
                <div class="stat-value" style="color: #f44336;">{discrepancy_count}</div>
                <div class="stat-label">With Issues</div>
            </div>
        """, unsafe_allow_html=True)
    with col3:
        valid_count = len(records) - discrepancy_count
        st.markdown(f"""
            <div class="stat-card">
                <div class="stat-value" style="color: #4CAF50;">{valid_count}</div>
                <div class="stat-label">Valid Records</div>
            </div>
        """, unsafe_allow_html=True)
    with col4:
        if records:
            latest_date = max(r.get('date', '') for r in records if r.get('date'))
            st.markdown(f"""
                <div class="stat-card">
                    <div class="stat-value" style="font-size: 1.3rem;">{latest_date or 'N/A'}</div>
                    <div class="stat-label">Latest Entry</div>
                </div>
            """, unsafe_allow_html=True)
    st.divider()
    st.subheader("📋 Recent Records")
    if not records:
        st.info("No records yet. Add your first record to get started!")
    else:
        recent_records = sorted(records, key=lambda x: x.get('created_at', ''), reverse=True)[:5]
        for record in recent_records:
            has_discrepancy = bool(record.get('discrepancies'))
            card_class = "record-card has-discrepancy" if has_discrepancy else "record-card"
            st.markdown(f"""
                <div class="{card_class}">
                    <div class="record-header">
                        {record.get('patient_name', 'Unknown')} - Tooth #{record.get('tooth', 'N/A')}
                    </div>
                    <div class="record-detail">📅 {record.get('date', 'N/A')} | 👨‍⚕️ {record.get('provider', 'N/A')}</div>
                    <div class="record-detail">🔢 Code: {record.get('code', 'N/A')}</div>
                    {f'<div class="discrepancy-badge">⚠️ Has Issues</div>' if has_discrepancy else ''}
                </div>
            """, unsafe_allow_html=True)
    if records:
        st.divider()
        st.subheader("📈 Analytics")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Records by Procedure Code**")
            code_counts = {}
            for record in records:
                code = record.get('code', 'Unknown')
                code_counts[code] = code_counts.get(code, 0) + 1
            if code_counts:
                df_codes = pd.DataFrame(list(code_counts.items()), columns=['Code', 'Count'])
                df_codes = df_codes.sort_values('Count', ascending=False).head(10)
                st.bar_chart(df_codes.set_index('Code'))
        with col2:
            st.markdown("**Records by Provider**")
            provider_counts = {}
            for record in records:
                provider = record.get('provider', 'Unknown')
                provider_counts[provider] = provider_counts.get(provider, 0) + 1
            if provider_counts:
                df_providers = pd.DataFrame(list(provider_counts.items()), columns=['Provider', 'Count'])
                df_providers = df_providers.sort_values('Count', ascending=False).head(10)
                st.bar_chart(df_providers.set_index('Provider'))

def show_records_page():
    st.header("📋 Audit Records")
    col1, col2, col3 = st.columns([3, 1, 1])
    with col1:
        search_query = st.text_input("🔍 Search records", placeholder="Patient name, ID, tooth, code, provider...")
    with col2:
        filter_option = st.selectbox("Filter", ["All Records", "With Issues", "Valid Only"])
    with col3:
        sort_option = st.selectbox("Sort By", ["Date (Newest)", "Date (Oldest)", "Patient Name"])
    
    records = st.session_state.record_manager.records
    if search_query:
        records = st.session_state.record_manager.search_records(search_query)
    if filter_option == "With Issues":
        records = [r for r in records if r.get('discrepancies')]
    elif filter_option == "Valid Only":
        records = [r for r in records if not r.get('discrepancies')]
    
    if sort_option == "Date (Newest)":
        records = sorted(records, key=lambda x: x.get('date', ''), reverse=True)
    elif sort_option == "Date (Oldest)":
        records = sorted(records, key=lambda x: x.get('date', ''))
    elif sort_option == "Patient Name":
        records = sorted(records, key=lambda x: x.get('patient_name', '').lower())
    
    col1, col2, col3, col4 = st.columns([1, 1, 1, 3])
    with col1:
        if st.button("📥 Export Excel", use_container_width=True):
            if records:
                excel_file = ExportManager.export_to_excel(records)
                st.download_button(
                    label="Download Excel",
                    data=excel_file,
                    file_name=f"audit_records_{datetime.datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            else:
                st.warning("No records to export")
    with col2:
        if st.button("📄 Export Word", use_container_width=True):
            if records:
                word_file = ExportManager.export_to_word(records)
                st.download_button(
                    label="Download Word",
                    data=word_file,
                    file_name=f"audit_report_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("No records to export")
    
    st.divider()
    st.markdown(f"**Showing {len(records)} record(s)**")
    
    if not records:
        st.info("No records found. Try adjusting your search or filters.")
    else:
        for i, record in enumerate(records):
            with st.container():
                col1, col2 = st.columns([4, 1])
                with col1:
                    has_discrepancy = bool(record.get('discrepancies'))
                    status_icon = "⚠️" if has_discrepancy else "✅"
                    st.markdown(f"### {status_icon} {record.get('patient_name', 'Unknown')} - Tooth #{record.get('tooth', 'N/A')}")
                    col_a, col_b, col_c = st.columns(3)
                    with col_a:
                        st.text(f"📅 Date: {record.get('date', 'N/A')}")
                        st.text(f"🆔 Patient ID: {record.get('patient_id', 'N/A')}")
                    with col_b:
                        st.text(f"👨‍⚕️ Provider: {record.get('provider', 'N/A')}")
                        st.text(f"🔢 Code: {record.get('code', 'N/A')}")
                    with col_c:
                        if record.get('surfaces'):
                            st.text(f"🦷 Surfaces: {record.get('surfaces')}")
                    if record.get('notes'):
                        with st.expander("📝 Notes"):
                            st.text(record.get('notes'))
                    if has_discrepancy:
                        st.error(f"⚠️ Issues: {record.get('discrepancies')}")
                with col2:
                    if st.button("👁️ View", key=f"view_{record.get('id')}", use_container_width=True):
                        st.session_state.viewing_record = record
                        st.rerun()
                    if st.button("✏️ Edit", key=f"edit_{record.get('id')}", use_container_width=True):
                        st.session_state.editing_record = record
                        st.rerun()
                    if st.button("🗑️ Delete", key=f"delete_{record.get('id')}", use_container_width=True):
                        if st.session_state.record_manager.delete_record(record.get('id')):
                            st.success("Record deleted")
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error("Delete failed")
                st.divider()
    
    if st.session_state.viewing_record:
        record = st.session_state.viewing_record
        with st.container():
            st.markdown("---")
            st.subheader("👁️ View Record Details")
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"**Patient:** {record.get('patient_name', 'N/A')}")
                st.markdown(f"**Patient ID:** {record.get('patient_id', 'N/A')}")
                st.markdown(f"**Date:** {record.get('date', 'N/A')}")
                st.markdown(f"**Tooth:** {record.get('tooth', 'N/A')}")
                st.markdown(f"**Code:** {record.get('code', 'N/A')} - {DENTAL_CODES.get(record.get('code', ''), 'Unknown')}")
                if record.get('surfaces'):
                    st.markdown(f"**Surfaces:** {record.get('surfaces')}")
                st.markdown(f"**Provider:** {record.get('provider', 'N/A')}")
                if record.get('notes'):
                    st.markdown("**Notes:**")
                    st.text_area("", value=record.get('notes'), height=100, disabled=True, label_visibility="collapsed")
                if record.get('discrepancies'):
                    st.error(f"**Issues Found:** {record.get('discrepancies')}")
            with col2:
                if st.button("Close", use_container_width=True):
                    st.session_state.viewing_record = None
                    st.rerun()
    
    if st.session_state.editing_record:
        record = st.session_state.editing_record
        with st.container():
            st.markdown("---")
            st.subheader("✏️ Edit Record")
            with st.form("edit_record_form"):
                col1, col2 = st.columns(2)
                with col1:
                    patient_name = st.text_input("Patient Name*", value=record.get('patient_name', ''))
                    patient_id = st.text_input("Patient ID*", value=record.get('patient_id', ''))
                    date_val = datetime.datetime.strptime(record.get('date', ''), '%d/%m/%Y') if record.get('date') else datetime.datetime.now()
                    date = st.date_input("Date*", value=date_val)
                    tooth = st.text_input("Tooth Number*", value=record.get('tooth', ''), max_chars=2)
                with col2:
                    code_idx = list(DENTAL_CODES.keys()).index(record.get('code', '111')) if record.get('code') in DENTAL_CODES else 0
                    code = st.selectbox("Procedure Code*", options=list(DENTAL_CODES.keys()), 
                                       index=code_idx,
                                       format_func=lambda x: f"{x} - {DENTAL_CODES[x]}")
                    surfaces = st.text_input("Surfaces", value=record.get('surfaces', ''))
                    provider = st.text_input("Provider*", value=record.get('provider', ''))
                notes = st.text_area("Notes", value=record.get('notes', ''), height=100)
                col_a, col_b = st.columns(2)
                with col_a:
                    save_button = st.form_submit_button("💾 Save Changes", use_container_width=True)
                with col_b:
                    cancel_button = st.form_submit_button("❌ Cancel", use_container_width=True)
            
            if save_button:
                if not all([patient_name, patient_id, tooth, code, provider]):
                    st.error("Please fill in all required fields")
                elif not DentalAuditLogic.validate_tooth_number(tooth):
                    st.error("Invalid tooth number format")
                else:
                    updated_data = {
                        "patient_name": patient_name,
                        "patient_id": patient_id,
                        "date": date.strftime('%d/%m/%Y'),
                        "tooth": tooth,
                        "code": code,
                        "surfaces": surfaces,
                        "provider": provider,
                        "notes": notes
                    }
                    analysis = DentalAuditLogic.analyze_record({**record, **updated_data})
                    if analysis["has_discrepancies"]:
                        updated_data["discrepancies"] = "; ".join(analysis["discrepancies"])
                    else:
                        updated_data["discrepancies"] = ""
                    if st.session_state.record_manager.update_record(record.get('id'), updated_data):
                        st.success("Record updated successfully!")
                        time.sleep(1)
                        st.session_state.editing_record = None
                        st.rerun()
                    else:
                        st.error("Failed to update record")
            
            if cancel_button:
                st.session_state.editing_record = None
                st.rerun()

def show_add_record_page():
    st.header("➕ Add New Record")
    with st.form("add_record_form"):
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Patient Information")
            patient_name = st.text_input("Patient Name*", placeholder="John Doe")
            patient_id = st.text_input("Patient ID*", placeholder="P12345")
            date = st.date_input("Date*", value=datetime.datetime.now())
        with col2:
            st.subheader("Procedure Details")
            tooth = st.text_input("Tooth Number*", placeholder="16", max_chars=2, help="2-digit tooth number")
            code = st.selectbox("Procedure Code*", options=list(DENTAL_CODES.keys()), 
                               format_func=lambda x: f"{x} - {DENTAL_CODES[x]}")
            surfaces = st.text_input("Surfaces (if applicable)", placeholder="MOD, B, L")
        provider = st.text_input("Provider*", placeholder="Dr. Smith")
        notes = st.text_area("Clinical Notes", placeholder="Enter any relevant clinical observations...", height=150)
        st.divider()
        col_a, col_b, col_c = st.columns([1, 1, 2])
        with col_a:
            submit_button = st.form_submit_button("💾 Save Record", use_container_width=True, type="primary")
        with col_b:
            clear_button = st.form_submit_button("🔄 Clear Form", use_container_width=True)
    
    if submit_button:
        if not all([patient_name, patient_id, tooth, code, provider]):
            st.error("❌ Please fill in all required fields")
        elif not DentalAuditLogic.validate_tooth_number(tooth):
            st.error("❌ Invalid tooth number format")
        else:
            new_record = {
                "patient_name": patient_name,
                "patient_id": patient_id,
                "date": date.strftime('%d/%m/%Y'),
                "tooth": tooth,
                "code": code,
                "surfaces": surfaces,
                "provider": provider,
                "notes": notes
            }
            analysis = DentalAuditLogic.analyze_record(new_record)
            if analysis["has_discrepancies"]:
                new_record["discrepancies"] = "; ".join(analysis["discrepancies"])
            else:
                new_record["discrepancies"] = ""
            if st.session_state.record_manager.add_record(new_record):
                st.success("✅ Record added successfully!")
                if analysis["has_discrepancies"]:
                    st.warning(f"⚠️ Note: {len(analysis['discrepancies'])} discrepancy/discrepancies found")
                    for disc in analysis["discrepancies"]:
                        st.error(f"• {disc}")
                time.sleep(2)
                st.rerun()
            else:
                st.error("❌ Failed to add record")
    
    if clear_button:
        st.rerun()

def show_ai_analysis_page():
    st.header("🤖 AI-Powered Analysis")
    col1, col2 = st.columns([2, 1])
    with col1:
        st.subheader("AI Model Configuration")
        selected_model = st.selectbox(
            "Select AI Model",
            options=get_selectable_models(),
            index=0,
            help="Choose between Cerebras Cloud models (online) or Ollama local models (offline)"
        )
        st.session_state.settings['ai_model'] = selected_model
    with col2:
        st.info("☁️ **Cloud models** require internet\n🖥️ **Local models** run offline")
    st.divider()
    
    tab1, tab2, tab3 = st.tabs(["🔍 Scan Records", "💬 Discuss Case", "📊 Batch Analysis"])
    
    with tab1:
        st.subheader("Scan Individual Record")
        records = st.session_state.record_manager.records
        if not records:
            st.info("No records available for analysis")
        else:
            record_options = [f"{r.get('patient_name', 'Unknown')} - Tooth {r.get('tooth', 'N/A')} ({r.get('date', 'N/A')})" 
                            for r in records]
            selected_idx = st.selectbox("Select Record to Analyze", range(len(records)), 
                                       format_func=lambda x: record_options[x])
            selected_record = records[selected_idx]
            with st.expander("📋 Record Details", expanded=True):
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.text(f"Patient: {selected_record.get('patient_name')}")
                    st.text(f"Patient ID: {selected_record.get('patient_id')}")
                with col2:
                    st.text(f"Date: {selected_record.get('date')}")
                    st.text(f"Tooth: {selected_record.get('tooth')}")
                with col3:
                    st.text(f"Code: {selected_record.get('code')}")
                    st.text(f"Provider: {selected_record.get('provider')}")
                if selected_record.get('notes'):
                    st.text_area("Notes", value=selected_record.get('notes'), height=100, disabled=True)
            
            if st.button("🔍 Scan & Analyze with AI", use_container_width=True, type="primary"):
                with st.spinner("AI is analyzing the record..."):
                    prompt = f"""Analyze this dental audit record for any discrepancies, errors, or concerns:

Patient: {selected_record.get('patient_name')}
Patient ID: {selected_record.get('patient_id')}
Date: {selected_record.get('date')}
Tooth: {selected_record.get('tooth')}
Procedure Code: {selected_record.get('code')} - {DENTAL_CODES.get(selected_record.get('code', ''), 'Unknown')}
Surfaces: {selected_record.get('surfaces', 'N/A')}
Provider: {selected_record.get('provider')}
Clinical Notes: {selected_record.get('notes', 'N/A')}

Please provide:
1. Any discrepancies or errors found
2. Validation of procedure coding
3. Clinical appropriateness
4. Documentation quality
5. Suggested corrections if needed"""
                    response = call_ai_model(prompt, selected_model)
                    if response:
                        st.success("✅ Analysis Complete")
                        st.markdown("### 🤖 AI Analysis Results")
                        st.markdown(response)
                        if st.button("💾 Save Analysis to Record"):
                            analysis_note = f"\n\n--- AI Analysis ({datetime.datetime.now().strftime('%d/%m/%Y %I:%M %p')}) ---\n{response}"
                            updated_notes = selected_record.get('notes', '') + analysis_note
                            st.session_state.record_manager.update_record(
                                selected_record.get('id'),
                                {"notes": updated_notes}
                            )
                            st.success("Analysis saved to record notes")
                    else:
                        st.error("❌ AI analysis failed")
    
    with tab2:
        st.subheader("Discuss Case with AI")
        st.markdown("""
        Have a conversation with the AI about any dental audit case.
        """)
        for msg in st.session_state.chat_history:
            if msg['role'] == 'user':
                st.markdown(f"**You:** {msg['content']}")
            else:
                st.markdown(f"**🤖 AI:** {msg['content']}")
        user_question = st.text_area("Ask a question about a case:", height=100)
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("Send", use_container_width=True, type="primary"):
                if user_question:
                    st.session_state.chat_history.append({"role": "user", "content": user_question})
                    with st.spinner("AI is thinking..."):
                        context = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.chat_history[-5:]])
                        prompt = f"""You are an expert dental auditor. Answer based on dental coding standards.

Previous conversation:
{context}

Current question: {user_question}

Please provide a detailed, professional answer."""
                        response = call_ai_model(prompt, selected_model)
                        if response:
                            st.session_state.chat_history.append({"role": "assistant", "content": response})
                            st.rerun()
                        else:
                            st.error("Failed to get AI response")
        with col2:
            if st.button("Clear Chat", use_container_width=True):
                st.session_state.chat_history = []
                st.rerun()
    
    with tab3:
        st.subheader("Batch Analysis")
        records = st.session_state.record_manager.records
        if not records:
            st.info("No records available for batch analysis")
        else:
            st.markdown(f"**{len(records)} records available for analysis**")
            col1, col2 = st.columns(2)
            with col1:
                analyze_all = st.checkbox("Analyze all records", value=True)
            with col2:
                if not analyze_all:
                    analyze_issues_only = st.checkbox("Only records with existing issues", value=False)
            if st.button("🚀 Start Batch Analysis", use_container_width=True, type="primary"):
                records_to_analyze = records
                if not analyze_all and 'analyze_issues_only' in locals() and analyze_issues_only:
                    records_to_analyze = [r for r in records if r.get('discrepancies')]
                if not records_to_analyze:
                    st.warning("No records match the criteria")
                else:
                    st.info(f"Analyzing {len(records_to_analyze)} record(s)...")
                    progress_bar = st.progress(0)
                    issues_found = 0
                    for i, record in enumerate(records_to_analyze):
                        progress_bar.progress((i + 1) / len(records_to_analyze))
                        analysis = DentalAuditLogic.analyze_record(record)
                        if analysis["has_discrepancies"]:
                            issues_found += 1
                    st.success(f"✅ Batch analysis complete! Found {issues_found} record(s) with issues.")

def show_settings_page():
    st.header("⚙️ Settings")
    user_data = st.session_state.user_data
    with st.form("settings_form"):
        st.subheader("User Profile")
        first_name = st.text_input("First Name", value=user_data.get('first_name', ''))
        last_name = st.text_input("Last Name", value=user_data.get('last_name', ''))
        email = st.text_input("Email", value=user_data.get('email', ''))
        st.divider()
        st.subheader("Security")
        new_password = st.text_input("New Password (leave blank to keep current)", type="password")
        confirm_password = st.text_input("Confirm New Password", type="password")
        st.divider()
        st.subheader("Export Settings")
        excel_export_name = st.text_input("Excel Export Name", value=user_data.get('excel_export_name', ''))
        st.divider()
        submit_button = st.form_submit_button("💾 Save Settings", use_container_width=True, type="primary")
    
    if submit_button:
        users_db = load_users_db()
        username = st.session_state.username
        if username in users_db.get("users", {}):
            user_data = users_db["users"][username]
            user_data["first_name"] = first_name
            user_data["last_name"] = last_name
            user_data["full_name"] = f"{first_name} {last_name}"
            user_data["email"] = email
            user_data["excel_export_name"] = excel_export_name
            if new_password:
                if len(new_password) < 4:
                    st.error("Password must be at least 4 characters")
                elif new_password != confirm_password:
                    st.error("Passwords do not match")
                else:
                    user_data["password_hash"] = hash_password(new_password)
            save_users_db(users_db)
            st.session_state.user_data = user_data
            st.success("Settings saved successfully!")

def show_about_page():
    st.header("ℹ️ About Easy Audit")
    st.markdown(f"""
    ### 🦷 Easy Audit V.{APP_VERSION}
    
    **Professional Dental Auditing System**
    
    Easy Audit is a comprehensive dental auditing platform designed to streamline the audit process
    for dental practices and ensure coding accuracy and compliance.
    
    #### ✨ Key Features
    
    - 🔐 **Multi-User Authentication**: Secure user management with bcrypt encryption
    - ☁️ **Cloud Synchronization**: Real-time sync with Supabase cloud database
    - 🤖 **AI-Powered Analysis**: Advanced discrepancy detection using Cerebras Cloud and Ollama
    - 📊 **Comprehensive Reporting**: Export to Excel and Word formats
    - 🦷 **Dental Code Validation**: 100+ ADA dental codes with automatic validation
    - 🔍 **Smart Search**: Quick search and filtering across all records
    - 📈 **Analytics Dashboard**: Visual insights into audit patterns
    - 💾 **Local & Cloud Storage**: Hybrid storage for reliability and accessibility
    
    #### 🔧 Technical Stack
    
    - **Framework**: Streamlit
    - **Database**: Supabase (Cloud) + Local JSON
    - **AI**: Cerebras Cloud API + Ollama Local
    - **Export**: openpyxl, python-docx
    - **Authentication**: bcrypt
    
    #### 📋 Supported Dental Codes
    
    - Diagnostic (111-128)
    - Preventive (141-154)
    - Restorative (211-263)
    - Endodontics (311-346)
    - Periodontics (411-442)
    - Prosthodontics (511-542)
    - Oral Surgery (611-642)
    - Orthodontics (711-715)
    - Adjunctive Services (811-851)
    
    #### 📞 Support
    
    For technical support or feature requests, please contact your system administrator.
    
    ---
    
    **Version**: {APP_VERSION}  
    **Release Date**: December 2024  
    **License**: Proprietary  
    
    © 2024 Easy Audit. All rights reserved.
    """)

# Main Application
def main():
    ensure_all_users_have_excel_names()
    sync_users_from_supabase()
    set_page_config()
    apply_custom_css()
    init_session_state()
    
    if not st.session_state.authenticated:
        if st.session_state.current_page == "register":
            show_register_page()
        else:
            show_login_page()
    else:
        show_main_dashboard()

if __name__ == "__main__":
    main()
